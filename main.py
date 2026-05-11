from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from playwright.sync_api import sync_playwright
import os
import zipfile
import time
import re
from datetime import datetime, timedelta
import base64
import json
import requests
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
from json_repair import repair_json
from copy import deepcopy
import threading
import psycopg2

app = Flask(__name__)
CORS(app, origins=["https://compita.umbusk.com"])

TEMP_DIR = "/tmp/descargas"
CACHE_DIAS = 30


def limpiar_archivos_viejos():
    try:
        if not os.path.exists(TEMP_DIR):
            return
        ahora = time.time()
        archivos_borrados = 0
        for archivo in os.listdir(TEMP_DIR):
            ruta_completa = os.path.join(TEMP_DIR, archivo)
            edad_segundos = ahora - os.path.getmtime(ruta_completa)
            edad_dias = edad_segundos / (60 * 60 * 24)
            if edad_dias > CACHE_DIAS:
                try:
                    os.remove(ruta_completa)
                    archivos_borrados += 1
                    print(f"Borrado: {archivo} (edad: {edad_dias:.1f} dias)")
                except Exception as e:
                    print(f"Error borrando {archivo}: {str(e)}")
        if archivos_borrados > 0:
            print(f"Limpieza completada: {archivos_borrados} archivos borrados")
    except Exception as e:
        print(f"Error en limpieza automatica: {str(e)}")


def verificar_archivo_en_cache(referencia):
    try:
        nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)
        for archivo in os.listdir(TEMP_DIR):
            if archivo.startswith(nombre_seguro) and archivo.endswith('_documento.pdf'):
                ruta_completa = os.path.join(TEMP_DIR, archivo)
                edad_segundos = time.time() - os.path.getmtime(ruta_completa)
                edad_dias = edad_segundos / (60 * 60 * 24)
                if edad_dias <= CACHE_DIAS:
                    print(f"Archivo en cache encontrado (edad: {edad_dias:.1f} dias)")
                    return ruta_completa
        return None
    except Exception as e:
        print(f"Error verificando cache: {str(e)}")
        return None


def descargar_pliego(referencia, guardar_zip=False):
    os.makedirs(TEMP_DIR, exist_ok=True)
    nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(accept_downloads=True)
        context.set_default_timeout(120000)
        page = context.new_page()

        try:
            print(f"Buscando licitacion: {referencia}")
            print(f"Navegando al portal...")
            url_listado = "https://comunidad.comprasdominicana.gob.do/Public/Tendering/ContractNoticeManagement/Index"
            page.goto(url_listado, timeout=90000)
            page.wait_for_timeout(5000)
            print(f"Portal cargado")

            print(f"Escribiendo en el buscador...")
            campo_busqueda = page.locator('#txtAllWords2Search')
            campo_busqueda.wait_for(state='visible', timeout=10000)
            campo_busqueda.clear()
            campo_busqueda.fill(referencia)
            print(f"Referencia ingresada: {referencia}")

            print(f"Buscando boton Buscar...")
            boton_encontrado = False

            try:
                boton_buscar = page.locator('input[type="button"][value="Buscar"]').first
                if boton_buscar.is_visible(timeout=5000):
                    boton_buscar.click()
                    boton_encontrado = True
                    print(f"Clic en boton azul (Opcion 1)")
            except:
                pass

            if not boton_encontrado:
                try:
                    boton_buscar = page.locator('input[value="Buscar"]').first
                    if boton_buscar.is_visible(timeout=5000):
                        boton_buscar.click()
                        boton_encontrado = True
                        print(f"Clic en boton (Opcion 2)")
                except:
                    pass

            if not boton_encontrado:
                try:
                    campo_busqueda.press('Enter')
                    boton_encontrado = True
                    print(f"Enter en campo de busqueda (Opcion 3)")
                except:
                    pass

            if not boton_encontrado:
                raise Exception("No se pudo ejecutar la busqueda - boton no encontrado")

            print(f"Esperando confirmacion del filtro...")
            page.wait_for_timeout(2000)
            filtro_aplicado = False

            try:
                indicador_filtro = page.locator('text=Buscar resultados por').first
                if indicador_filtro.is_visible(timeout=20000):
                    print(f"Filtro confirmado: Buscar resultados por")
                    filtro_aplicado = True
            except:
                pass

            if not filtro_aplicado:
                try:
                    link_borrar = page.locator('a:has-text("Borrar")').first
                    if link_borrar.is_visible(timeout=5000):
                        print(f"Filtro confirmado: Link Borrar busqueda")
                        filtro_aplicado = True
                except:
                    pass

            if filtro_aplicado:
                print(f"FILTRO APLICADO CORRECTAMENTE")
            else:
                print(f"ADVERTENCIA: No se confirmo el filtro, continuando...")

            page.wait_for_timeout(3000)

            print(f"Buscando {referencia} en resultados...")
            resultado = None
            xpaths = [
                f'//td[contains(text(), "{referencia}")]',
                f'//td[text()="{referencia}"]',
                f'//*[contains(text(), "{referencia}")]',
                f'//td[normalize-space(text())="{referencia}"]'
            ]

            for i, xpath in enumerate(xpaths):
                try:
                    print(f"   Intentando XPath {i+1}...")
                    resultado_xpath = page.locator(f'xpath={xpath}').first
                    if resultado_xpath.is_visible(timeout=10000):
                        resultado = resultado_xpath
                        print(f"   XPath {i+1} funciono")
                        break
                except:
                    print(f"   XPath {i+1} fallo")
                    continue

            if not resultado:
                raise Exception(f"No se encontro la licitacion {referencia} en los resultados")

            print(f"Licitacion encontrada en la tabla")

            screenshot_path = f"{TEMP_DIR}/debug_tabla.png"
            page.screenshot(path=screenshot_path)

            print(f"Buscando boton DETALLE...")
            fila = resultado.locator('xpath=ancestor::tr')

            boton_detalle = None
            selectores_detalle = [
                'a[title="Detalle"]',
                'a[id*="lnkDetailLink"]',
                'xpath=.//a[@title="Detalle"]',
                'xpath=.//a[text()="Detalle"]',
                'a[href="javascript:void(0)"]',
                'a:has-text("Detalle")',
                '*:has-text("Detalle")'
            ]

            for i, selector in enumerate(selectores_detalle):
                try:
                    print(f"   Probando selector {i+1}: {selector}")
                    boton = fila.locator(selector).first
                    if boton.count() > 0:
                        boton_detalle = boton
                        print(f"   Selector {i+1} funciono")
                        break
                    else:
                        print(f"   Selector {i+1} no encontro elementos")
                except Exception as e:
                    print(f"   Selector {i+1} error: {str(e)[:50]}")
                    continue

            if not boton_detalle:
                raise Exception("No se encontro el boton DETALLE con ningun selector")

            print(f"Haciendo scroll al boton...")
            boton_detalle.scroll_into_view_if_needed()
            page.wait_for_timeout(2000)

            print(f"Haciendo clic en DETALLE...")
            boton_detalle.click()
            print(f"Clic en DETALLE realizado")

            page.wait_for_timeout(5000)

            pages = context.pages
            if len(pages) > 1:
                page = pages[-1]
                print(f"Cambiado a nueva ventana")

            print(f"Buscando iframe del detalle...")
            page.wait_for_timeout(3000)

            frames = page.frames
            print(f"   Total frames: {len(frames)}")

            iframe_correcto = None
            for i, frame in enumerate(frames):
                try:
                    print(f"   Probando frame {i+1}...")
                    boton_test = frame.locator('#tbToolBar_btnTbDownload').first
                    if boton_test.count() > 0:
                        print(f"   FRAME CORRECTO encontrado")
                        iframe_correcto = frame
                        break
                    else:
                        print(f"   Frame no tiene el boton")
                except Exception as e:
                    print(f"   Error en frame {i+1}: {str(e)[:80]}")
                    continue

            if not iframe_correcto:
                print(f"   Reintentando busqueda por contenido de texto...")
                for i, frame in enumerate(frames):
                    try:
                        body_text = frame.locator('body').text_content(timeout=5000)
                        if body_text and referencia in body_text:
                            print(f"   Frame {i+1} contiene la referencia")
                            iframe_correcto = frame
                            break
                    except:
                        continue

            if not iframe_correcto:
                raise Exception("No se encontro iframe con el boton de descarga")

            print(f"Buscando boton de descarga...")
            boton_descarga = None
            selectores_descarga = [
                '#tbToolBar_btnTbDownload',
                'input[id="tbToolBar_btnTbDownload"]',
                'input[title="Descargar procedimiento"]',
                'input[value="Descargar procedimiento"]',
                'input[type="button"][value="Descargar procedimiento"]'
            ]

            for i, selector in enumerate(selectores_descarga):
                try:
                    print(f"   Probando selector {i+1}: {selector}")
                    boton = iframe_correcto.locator(selector).first
                    if boton.count() > 0:
                        boton_descarga = boton
                        print(f"   Selector {i+1} funciono")
                        break
                    else:
                        print(f"   Selector {i+1} no encontro elementos")
                except Exception as e:
                    print(f"   Selector {i+1} error: {str(e)[:50]}")
                    continue

            if not boton_descarga:
                raise Exception("No se encontro el boton de descarga")

            print(f"Iniciando descarga...")
            with page.expect_download(timeout=90000) as download_info:
                boton_descarga.click()

            download = download_info.value
            zip_path = f"{TEMP_DIR}/{nombre_seguro}.zip"
            download.save_as(zip_path)
            print(f"ZIP descargado: {zip_path}")

            if not os.path.exists(zip_path):
                raise Exception("El archivo ZIP no se guardo correctamente")

            tamano_mb = os.path.getsize(zip_path) / (1024 * 1024)
            print(f"   Tamano: {tamano_mb:.2f} MB")

            print(f"Extrayendo documento principal...")
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                archivos_zip = zip_ref.namelist()
                print(f"   Archivos en ZIP: {len(archivos_zip)}")
                documento_encontrado = None
                palabras_clave = ['pliego', 'ficha', 'terminos', 'especificaciones', 'anexo']

                for palabra in palabras_clave:
                    if documento_encontrado:
                        break
                    for archivo in archivos_zip:
                        if '1_Publicaciones/Adjuntos/' in archivo:
                            nombre_archivo = os.path.basename(archivo).lower()
                            if palabra in nombre_archivo and archivo.endswith('.pdf'):
                                print(f"   Documento encontrado por '{palabra}': {os.path.basename(archivo)}")
                                documento_path = f"{TEMP_DIR}/{nombre_seguro}_{int(time.time())}_documento.pdf"
                                with zip_ref.open(archivo) as source:
                                    with open(documento_path, 'wb') as target:
                                        target.write(source.read())
                                documento_encontrado = documento_path
                                break

                if not documento_encontrado:
                    print(f"   No se encontro por palabras clave, buscando PDF mas grande...")
                    pdfs_en_adjuntos = []
                    for archivo in archivos_zip:
                        if '1_Publicaciones/Adjuntos/' in archivo and archivo.endswith('.pdf'):
                            info = zip_ref.getinfo(archivo)
                            pdfs_en_adjuntos.append({'nombre': archivo, 'tamano': info.file_size})

                    if pdfs_en_adjuntos:
                        pdfs_en_adjuntos.sort(key=lambda x: x['tamano'], reverse=True)
                        archivo_mas_grande = pdfs_en_adjuntos[0]['nombre']
                        print(f"   Usando PDF mas grande: {os.path.basename(archivo_mas_grande)}")
                        documento_path = f"{TEMP_DIR}/{nombre_seguro}_{int(time.time())}_documento.pdf"
                        with zip_ref.open(archivo_mas_grande) as source:
                            with open(documento_path, 'wb') as target:
                                target.write(source.read())
                        documento_encontrado = documento_path

                if not documento_encontrado:
                    raise Exception("No se encontro ningun documento principal en el ZIP")

                print(f"Documento extraido exitosamente")

                if not guardar_zip:
                    try:
                        if os.path.exists(zip_path):
                            os.remove(zip_path)
                            print(f"ZIP eliminado")
                    except Exception as e:
                        print(f"Error borrando ZIP: {str(e)}")
                else:
                    print(f"ZIP conservado: {zip_path}")

                browser.close()
                return documento_encontrado

        except Exception as e:
            browser.close()
            raise Exception(f"Error al descargar: {str(e)}")


@app.route('/descargar-pliego', methods=['POST'])
def endpoint_descargar_pliego():
    try:
        data = request.get_json()
        referencia = data.get('referencia')

        if not referencia:
            return jsonify({"error": "Falta el parametro 'referencia'"}), 400

        limpiar_archivos_viejos()
        documento_path = verificar_archivo_en_cache(referencia)

        if documento_path:
            print(f"Usando documento en cache")
        else:
            print(f"Descargando documento (no esta en cache)")
            documento_path = descargar_pliego(referencia)

        nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)

        return send_file(
            documento_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"documento_{nombre_seguro}.pdf"
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


def buscar_precios_referencia(titulo, descripcion):
    db_url = os.environ.get('DATABASE_URL')
    if not db_url:
        return []
    try:
        terminos = f"{titulo} {descripcion}".strip()[:300]
        conn = psycopg2.connect(db_url)
        cur = conn.cursor()
        cur.execute("""
            SELECT
                descripcion,
                unidad_medida,
                ROUND(AVG(precio_unitario)::numeric, 2)  AS precio_promedio,
                ROUND(MIN(precio_unitario)::numeric, 2)  AS precio_minimo,
                ROUND(MAX(precio_unitario)::numeric, 2)  AS precio_maximo,
                COUNT(*)                                  AS num_referencias,
                moneda
            FROM precios_referencia
            WHERE to_tsvector('spanish', descripcion)
                  @@ plainto_tsquery('spanish', %s)
              AND precio_unitario > 0
            GROUP BY descripcion, unidad_medida, moneda
            ORDER BY num_referencias DESC
            LIMIT 10
        """, (terminos,))
        filas = cur.fetchall()
        cur.close()
        conn.close()
        return [
            {
                "descripcion":     f[0],
                "unidad_medida":   f[1],
                "precio_promedio": float(f[2]) if f[2] else None,
                "precio_minimo":   float(f[3]) if f[3] else None,
                "precio_maximo":   float(f[4]) if f[4] else None,
                "num_referencias": int(f[5]),
                "moneda":          f[6] or "DOP"
            }
            for f in filas
        ]
    except Exception as e:
        print(f"Error buscando precios referencia: {e}")
        return []


@app.route('/analizar-pliego', methods=['POST'])
def analizar_pliego():
    try:
        data = request.get_json()
        referencia = data.get('referencia')
        titulo = data.get('titulo', '')
        descripcion = data.get('descripcion', '')
        monto = data.get('monto', 0)
        empresa_descripcion = data.get('empresa_descripcion', '')
        empresa_website = data.get('empresa_website', '')
        fecha_presentacion = data.get('fecha_presentacion', '')
        fecha_hoy = datetime.now().strftime('%d/%m/%Y')

        if not referencia:
            return jsonify({'success': False, 'error': 'Falta referencia'}), 400

        print(f"Iniciando analisis de {referencia}")

        archivo_pdf = verificar_archivo_en_cache(referencia)

        if not archivo_pdf:
            print(f"Archivo no encontrado en cache, descargando...")
            archivo_pdf = descargar_pliego(referencia)
        else:
            print(f"Usando archivo desde cache: {archivo_pdf}")

        tamano_bytes = os.path.getsize(archivo_pdf)
        tamano_mb = tamano_bytes / (1024 * 1024)
        print(f"Tamano del PDF: {tamano_mb:.2f} MB")

        if tamano_mb > 50:
            return jsonify({
                'success': False,
                'error': f'El pliego es demasiado grande ({tamano_mb:.1f} MB). Limite: 50 MB.'
            }), 400

        print(f"Extrayendo texto del PDF...")
        es_pdf_imagen = False
        texto_completo = ""
        try:
            reader = PdfReader(archivo_pdf)

            for i, page in enumerate(reader.pages):
                try:
                    texto_pagina = page.extract_text()
                    if texto_pagina:
                        texto_completo += texto_pagina + "\n\n"
                except Exception as e:
                    print(f"Error extrayendo pagina {i+1}: {str(e)}")
                    continue

            if len(texto_completo.strip()) < 200:
                print(f"Texto insuficiente ({len(texto_completo)} chars) — PDF es imagen, usando envio nativo a Claude")
                es_pdf_imagen = True
                texto_completo = "[El pliego es un PDF de imagen — se adjunta como documento para analisis directo por Claude]"
            else:
                if len(texto_completo) > 100000:
                    print(f"Texto muy largo ({len(texto_completo)} caracteres), truncando...")
                    texto_completo = texto_completo[:100000] + "\n\n[DOCUMENTO TRUNCADO]"
                print(f"Texto extraido: {len(texto_completo)} caracteres, {len(reader.pages)} paginas")

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Error al procesar el PDF: {str(e)}'
            }), 400

        print(f"Buscando precios historicos de referencia...")
        precios_ref = buscar_precios_referencia(titulo, descripcion)
        print(f"Precios historicos encontrados: {len(precios_ref)} items")

        bloque_precios = ""
        if precios_ref:
            lineas = []
            for p in precios_ref:
                lineas.append(
                    f"- {p['descripcion']} ({p['unidad_medida'] or 'sin unidad'}): "
                    f"promedio {p['precio_promedio']:,.2f} {p['moneda']} "
                    f"[rango {p['precio_minimo']:,.2f} – {p['precio_maximo']:,.2f}] "
                    f"({p['num_referencias']} ofertas historicas)"
                )
            bloque_precios = f"""
PRECIOS HISTORICOS DE REFERENCIA (base de datos interna Compita, licitaciones adjudicadas 2026):
{chr(10).join(lineas)}

Usa estos precios como contexto para evaluar si los montos estimados del pliego son razonables.
Si algun item del pliego coincide con los anteriores, menciónalo en tu analisis.
"""

        perfil_empresa = ""
        if empresa_descripcion:
            perfil_empresa += f"\n- Descripcion: {empresa_descripcion}"
        if empresa_website:
            perfil_empresa += f"\n- Sitio web: {empresa_website}"

        seccion_perfil = f"""
PERFIL DE LA EMPRESA QUE EVALUA:
{perfil_empresa if perfil_empresa else "No disponible"}
""" if perfil_empresa else ""

        prompt_analisis = f"""Eres un experto analista de licitaciones publicas dominicanas.

CONTEXTO DE LA LICITACION:
- Referencia: {referencia}
- Titulo: {titulo}
- Descripcion: {descripcion}
- Monto estimado: RD${monto:,.2f}
- Fecha de hoy: {fecha_hoy}
- Fecha limite de presentacion: {fecha_presentacion if fecha_presentacion else 'No disponible'}
{seccion_perfil}{bloque_precios}
A continuacion esta el contenido completo del pliego de condiciones:

---INICIO DEL PLIEGO---
{texto_completo}
---FIN DEL PLIEGO---

INSTRUCCIONES:
Analiza el pliego y proporciona un analisis estructurado en formato JSON con esta estructura exacta:

{{
  "sintesis": "Resumen ejecutivo en 2-3 oraciones",
  "oportunidades": ["Primera oportunidad", "Segunda oportunidad", "Tercera oportunidad"],
  "riesgos": ["Primer riesgo", "Segundo riesgo", "Tercer riesgo"],
  "requisitos": ["Primer requisito", "Segundo requisito", "Tercer requisito"],
  "certificaciones_iso": {{
    "exige_iso": "SI o NO",
    "listado": ["ISO XXXX - descripcion"],
    "nota": "Normas tecnicas equivalentes si aplica"
  }},
  "tiempos": {{
    "fecha_limite_oferta": "DD/MM/YYYY",
    "dias_calendario_restantes": "N dias desde hoy ({fecha_hoy})",
    "alerta": "HOLGADO, AJUSTADO, o MUY AJUSTADO",
    "fechas_clave": ["Lista de fechas relevantes del pliego"],
    "advertencia": "Si tiempo ajustado, explicar impacto. Vacio si holgado."
  }},
  "viabilidad": {{
    "veredicto": "VIABLE, VIABLE CON RIESGOS, o DIFICIL DE CUMPLIR",
    "garantias": "Descripcion de garantias exigidas",
    "experiencia_previa": "Experiencia previa que exige el pliego",
    "especificaciones_tecnicas": "Compatibilidad con perfil de la empresa"
  }},
  "evaluacion": {{
    "a_favor": ["Argumento 1", "Argumento 2", "Argumento 3"],
    "en_contra": ["Riesgo 1", "Riesgo 2", "Riesgo 3"]
  }},
  "precios_historicos": {json.dumps([
      {"item": p["descripcion"][:80],
       "precio_promedio": p["precio_promedio"],
       "precio_minimo": p["precio_minimo"],
       "precio_maximo": p["precio_maximo"],
       "num_referencias": p["num_referencias"],
       "moneda": p["moneda"]}
      for p in precios_ref
  ], ensure_ascii=False) if precios_ref else "[]"}
}}

IMPORTANTE: Responde SOLO con el JSON, sin texto adicional ni markdown."""

        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            return jsonify({
                'success': False,
                'error': 'API key de Anthropic no configurada'
            }), 500

        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01"
        }

        if es_pdf_imagen:
            if tamano_mb > 20:
                return jsonify({
                    'success': False,
                    'error': f'El pliego ({tamano_mb:.1f} MB) no tiene texto extraible y supera el limite para analisis de imagen (~20 MB). Descargalo manualmente desde el portal SECP para revisarlo.'
                }), 400
            print("Enviando PDF imagen directamente a Claude AI (modo documento nativo)...")
            with open(archivo_pdf, 'rb') as f:
                pdf_b64 = base64.b64encode(f.read()).decode('utf-8')
            payload = {
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 3000,
                "messages": [{
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_b64
                            }
                        },
                        {
                            "type": "text",
                            "text": prompt_analisis
                        }
                    ]
                }]
            }
        else:
            print("Enviando texto del pliego a Claude AI...")
            payload = {
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 3000,
                "messages": [{"role": "user", "content": prompt_analisis}]
            }

        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers,
            json=payload,
            timeout=120
        )

        if response.status_code == 413:
            return jsonify({
                'success': False,
                'error': f'El pliego es demasiado grande para el analisis automatico ({tamano_mb:.1f} MB). Descargalo manualmente desde el portal SECP para revisarlo.'
            }), 400

        if response.status_code != 200:
            raise Exception(f"Error de Claude API: {response.status_code}")

        claude_response = response.json()
        analisis_texto = claude_response['content'][0]['text']
        analisis_texto = analisis_texto.replace('```json', '').replace('```', '').strip()

        inicio = analisis_texto.find('{')
        fin = analisis_texto.rfind('}')
        if inicio == -1 or fin == -1:
            raise json.JSONDecodeError('No se encontro JSON valido', analisis_texto, 0)
        analisis_texto = analisis_texto[inicio:fin+1]

        analisis = json.loads(analisis_texto)

        if 'precios_historicos' not in analisis:
            analisis['precios_historicos'] = precios_ref

        print(f"Analisis completado")

        return jsonify({
            'success': True,
            'pliego_analizado': True,
            'analisis': analisis,
            'tiene_precios_historicos': len(precios_ref) > 0
        })

    except json.JSONDecodeError as e:
        print(f"Error parseando JSON del analisis: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Error al procesar respuesta de Claude AI'
        }), 500

    except Exception as e:
        print(f"Error en analisis: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "service": "compita-descarga-pliegos"})


@app.route('/cache/info', methods=['GET'])
def cache_info():
    try:
        if not os.path.exists(TEMP_DIR):
            return jsonify({"archivos": 0, "tamano_total_mb": 0, "archivos_list": []})

        archivos_info = []
        tamano_total = 0

        for archivo in os.listdir(TEMP_DIR):
            ruta_completa = os.path.join(TEMP_DIR, archivo)
            tamano = os.path.getsize(ruta_completa)
            edad_segundos = time.time() - os.path.getmtime(ruta_completa)
            edad_dias = edad_segundos / (60 * 60 * 24)
            archivos_info.append({
                "nombre": archivo,
                "tamano_mb": round(tamano / (1024 * 1024), 2),
                "edad_dias": round(edad_dias, 1)
            })
            tamano_total += tamano

        return jsonify({
            "archivos": len(archivos_info),
            "tamano_total_mb": round(tamano_total / (1024 * 1024), 2),
            "cache_dias": CACHE_DIAS,
            "archivos_list": archivos_info
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/cache/limpiar', methods=['POST'])
def cache_limpiar():
    try:
        limpiar_archivos_viejos()
        return jsonify({"status": "ok", "mensaje": "Limpieza ejecutada"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def extraer_items_con_claude(pdf_bytes_list, referencia):

    def extraer_de_un_pdf(pdf_bytes, indice, contexto_previo=""):
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))

            if reader.is_encrypted:
                print(f"PDF {indice + 1} esta protegido con contrasena - omitido")
                return []

            texto = ""
            paginas_sin_texto = 0
            for pg in reader.pages:
                try:
                    t = pg.extract_text()
                    if t and t.strip():
                        texto += t + "\n"
                    else:
                        paginas_sin_texto += 1
                except Exception as e:
                    print(f"Error extrayendo pagina en PDF {indice + 1}: {e}")
                    paginas_sin_texto += 1
                    continue

            total_paginas = len(reader.pages)
            es_imagen = False
            if total_paginas > 0 and paginas_sin_texto == total_paginas:
                print(f"PDF {indice + 1} es imagen ({total_paginas} pags sin texto) — enviando como documento nativo a Claude")
                es_imagen = True
                texto = "[PDF de imagen — se adjunta como documento nativo para extraccion de items]"
            elif paginas_sin_texto > 0:
                print(f"PDF {indice + 1}: {paginas_sin_texto} de {total_paginas} paginas sin texto")

        except Exception as e:
            msg = str(e).lower()
            if 'password' in msg or 'encrypt' in msg:
                print(f"PDF {indice + 1} requiere contrasena - omitido")
            elif 'eof' in msg or 'invalid' in msg or 'corrupt' in msg:
                print(f"PDF {indice + 1} parece estar corrupto - omitido")
            else:
                print(f"Error leyendo PDF {indice + 1}: {e}")
            return []

        if not texto.strip() and not es_imagen:
            print(f"PDF {indice + 1} no tiene texto extraible - omitido")
            return []

        if contexto_previo:
            contexto_str = f"""
CONTEXTO IMPORTANTE: Los documentos anteriores ya extrajeron los siguientes items:
{contexto_previo}
Este documento es una CONTINUACION. Si no ves un encabezado de lote al inicio,
asigna los items al lote donde corresponde segun la numeracion que continua.
"""
        else:
            contexto_str = ""

        prompt = f"""Eres un experto en licitaciones publicas dominicanas y tributacion del ITBIS.
{contexto_str}
Contenido de un documento de la licitacion {referencia} (documento {indice + 1}):

{texto[:120000]}

INSTRUCCION:
Extrae TODOS los items, productos, equipos o materiales que aparecen en ESTE documento.
Pueden estar organizados en LOTES (LOTE I, LOTE II, LOTE III, etc.).
IMPORTANTE: identifica correctamente el lote de cada item segun los encabezados
"LOTE- I", "LOTE- II", "LOTE- III" que aparecen en el documento.
No asumas que todos los items son del mismo lote.

Ademas, determina la politica de ITBIS de esta licitacion:
- "TRANSPARENTADO": el pliego pide que el ITBIS se declare por separado
- "INCLUIDO": el pliego indica que los precios deben incluir todos los impuestos
- "EXENTO": el pliego declara que esta contratacion esta exenta de ITBIS
- "NO_ESPECIFICADO": el pliego no menciona nada sobre ITBIS

Para cada item, determina si aplica ITBIS segun la ley dominicana:
- itbis_aplica: false si el item es medicamento, reactivo de laboratorio, equipo medico
  de la lista oficial, sangre o derivados, insumos medicos de la canasta basica, o si
  la politica es EXENTO. En todos los demas casos, itbis_aplica: true.

Para cada item devuelve:
- lote: numero romano del lote (ej: "I", "II", "III"). Si no hay lotes, usa "I".
- numero: numero del item dentro de su lote
- descripcion: descripcion completa (incluye marca y modelo si aparecen)
- unidad: unidad de medida (UD, SVC, PAQ, KG, etc.)
- cantidad: cantidad numerica (o null si no aparece)
- itbis_aplica: true o false segun las reglas anteriores

Si este documento no contiene ningun producto, equipo o material, devuelve lista vacia.

Responde UNICAMENTE con JSON valido, sin texto adicional:
{{
  "politica_itbis": "TRANSPARENTADO",
  "items": [
    {{"lote": "I", "numero": "1", "descripcion": "...", "unidad": "UD", "cantidad": 1, "itbis_aplica": true}},
    {{"lote": "II", "numero": "1", "descripcion": "...", "unidad": "UD", "cantidad": 1, "itbis_aplica": false}}
  ]
}}"""

        api_key = os.environ.get('ANTHROPIC_API_KEY')
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01"
        }
        payload = {
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 16000,
            "messages": [{
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": base64.b64encode(pdf_bytes).decode('utf-8')
                        }
                    },
                    {"type": "text", "text": prompt}
                ]
            }]
        } if es_imagen else {
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 16000,
            "messages": [{"role": "user", "content": prompt}]
        }

        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers,
            json=payload,
            timeout=300
        )
        if resp.status_code != 200:
            print(f"Error Claude API en PDF {indice + 1}: {resp.status_code}")
            return []

        texto_resp = resp.json()['content'][0]['text']
        texto_resp = texto_resp.replace('```json', '').replace('```', '').strip()
        inicio = texto_resp.find('{')
        fin = texto_resp.rfind('}')
        if inicio == -1 or fin == -1:
            return []
        json_raw = texto_resp[inicio:fin + 1]
        try:
            data = json.loads(json_raw)
        except json.JSONDecodeError:
            print(f"JSON malformado en PDF {indice + 1}, aplicando reparacion...")
            json_reparado = repair_json(json_raw)
            data = json.loads(json_reparado)

        politica = data.get('politica_itbis', 'NO_ESPECIFICADO')
        items_pdf = data.get('items', [])

        for item in items_pdf:
            item['politica_itbis'] = politica

        print(f"  Politica ITBIS detectada: {politica}")
        return items_pdf

    todos_items = {}
    contexto_previo = ""

    for i, pdf_bytes in enumerate(pdf_bytes_list):
        print(f"  Procesando PDF {i + 1}/{len(pdf_bytes_list)} con Claude...")
        items_pdf = extraer_de_un_pdf(pdf_bytes, i, contexto_previo)
        print(f"     -> {len(items_pdf)} items encontrados")

        for item in items_pdf:
            lote = str(item.get('lote', 'I')).strip().upper()
            num = str(item.get('numero', '')).strip()
            if not num:
                continue
            clave = f"{lote}-{num}"
            if clave not in todos_items or not todos_items[clave].get('descripcion'):
                todos_items[clave] = item

        if items_pdf:
            resumen = {}
            for item in items_pdf:
                lote = str(item.get('lote', 'I')).strip().upper()
                resumen.setdefault(lote, []).append(
                    int(item.get('numero', 0)) if str(item.get('numero', 0)).isdigit() else 0
                )
            partes = []
            for lote in sorted(resumen):
                nums = sorted(resumen[lote])
                partes.append(f"Lote {lote}: items {nums[0]} al {nums[-1]} ({len(nums)} items)")
            contexto_previo = "; ".join(partes)

    orden_lote = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5}

    def clave_orden(item):
        lote = str(item.get('lote', 'I')).strip().upper()
        try:
            num = int(item.get('numero', 0))
        except (ValueError, TypeError):
            num = 9999
        return (orden_lote.get(lote, 99), num)

    return sorted(todos_items.values(), key=clave_orden)


def extraer_docx_de_adjuntos(zf):
    resultado = []
    for archivo in zf.namelist():
        if '1_Publicaciones/Adjuntos/' not in archivo:
            continue
        nombre = os.path.basename(archivo).lower()
        if archivo.lower().endswith(('.docx', '.doc')):
            resultado.append((nombre, zf.read(archivo)))
        elif archivo.lower().endswith('.zip'):
            try:
                nested_bytes = zf.read(archivo)
                with zipfile.ZipFile(io.BytesIO(nested_bytes), 'r') as nzf:
                    for na in nzf.namelist():
                        nn = os.path.basename(na).lower()
                        if na.lower().endswith(('.docx', '.doc')):
                            resultado.append((nn, nzf.read(na)))
            except Exception as e:
                print(f"Error leyendo ZIP anidado {nombre}: {e}")
    return resultado
def generar_oferta_desde_ficha(items, nombre_ficha="ficha tecnica"):
    """
    Genera un Word de oferta economica cuando la licitacion no tiene F033.
    Pre-llena item, descripcion, unidad, cantidad e ITBIS.
    Deja precio unitario y total en blanco para que el usuario los complete.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run("FORMULARIO DE OFERTA ECONOMICA")
    run_titulo.bold = True
    run_titulo.font.size = Pt(13)

    p_nota = doc.add_paragraph()
    run_label = p_nota.add_run("NOTA: ")
    run_label.bold = True
    run_label.font.size = Pt(8)
    run_label.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    run_texto = p_nota.add_run(
        f"Este formulario fue generado automaticamente por Agente 033 de Compita "
        f"a partir de la ficha tecnica \"{nombre_ficha}\". "
        f"Verifique los campos contra el pliego original antes de presentar. "
        f"Complete los precios unitarios y totales antes de la presentacion."
    )
    run_texto.font.size = Pt(8)
    run_texto.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    doc.add_paragraph()

    politica_global = 'NO_ESPECIFICADO'
    for item in items:
        p = item.get('politica_itbis', '')
        if p:
            politica_global = p
            break
    print(f"  Politica ITBIS (oferta desde ficha): {politica_global}")

    tabla = doc.add_table(rows=1, cols=7)
    tabla.style = 'Table Grid'

    anchos = [0.45, 3.0, 0.7, 0.7, 1.0, 0.6, 1.0]
    cabeceras = ['No.', 'Descripcion', 'Unidad', 'Cantidad',
                 'Precio Unitario', 'ITBIS', 'Total']

    fila_cab = tabla.rows[0]
    for i, (cab, ancho) in enumerate(zip(cabeceras, anchos)):
        cel = fila_cab.cells[i]
        cel.width = Inches(ancho)
        p = cel.paragraphs[0]
        p.clear()
        run = p.add_run(cab)
        run.bold = True
        run.font.size = Pt(8)

    for item in items:
        fila = tabla.add_row()
        celdas = fila.cells

        lote = str(item.get('lote', '')).strip()
        num  = item.get('numero', '')
        etiqueta = f"L-{lote}-{num}" if lote else str(num)

        itbis_aplica = item.get('itbis_aplica', True)
        if politica_global == 'EXENTO':
            valor_itbis = '0%'
        elif politica_global == 'INCLUIDO':
            valor_itbis = 'Incluido'
        else:
            valor_itbis = '18%' if itbis_aplica else '0%'

        valores = [
            etiqueta,
            item.get('descripcion', ''),
            item.get('unidad', ''),
            str(item.get('cantidad', '')) if item.get('cantidad') is not None else '',
            '',
            valor_itbis,
            '',
        ]
        for i, val in enumerate(valores):
            p = celdas[i].paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.size = Pt(8)

    fila_total = tabla.add_row()
    cel_label = fila_total.cells[0]
    for j in range(1, 5):
        cel_label = cel_label.merge(fila_total.cells[j])
    p_total = cel_label.paragraphs[0]
    p_total.clear()
    run_total = p_total.add_run('VALOR TOTAL')
    run_total.bold = True
    run_total.font.size = Pt(8)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def llenar_f033(docx_bytes, items):
    doc = Document(io.BytesIO(docx_bytes))
    tabla = None
    for t in doc.tables:
        if len(t.columns) >= 6:
            tabla = t
            break

    if not tabla:
        raise Exception("No se encontro la tabla del F033 en el Word")

    filas_datos = []
    fila_total = None
    fila_template = None

    for row in tabla.rows:
        txt = row.cells[0].text.strip().lower()
        if any(k in txt for k in ['item', 'no.', 'descripci', 'unidad']):
            continue
        if 'valor total' in row.cells[0].text.lower():
            fila_total = row
            continue
        filas_datos.append(row)
        fila_template = row

    if len(filas_datos) < len(items):
        if fila_total and fila_template:
            faltan = len(items) - len(filas_datos)
            for _ in range(faltan):
                nueva_tr = deepcopy(fila_template._tr)
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                for tc in nueva_tr.findall(f'.//{ns}tc'):
                    for p in tc.findall(f'{ns}p'):
                        for r in p.findall(f'{ns}r'):
                            p.remove(r)
                fila_total._tr.addprevious(nueva_tr)
            filas_datos = []
            for row in tabla.rows:
                txt = row.cells[0].text.strip().lower()
                if any(k in txt for k in ['item', 'no.', 'descripci', 'unidad']):
                    continue
                if 'valor total' in row.cells[0].text.lower():
                    break
                filas_datos.append(row)
        else:
            while len(filas_datos) < len(items):
                filas_datos.append(tabla.add_row())

    politica_global = 'NO_ESPECIFICADO'
    for item in items:
        p = item.get('politica_itbis', '')
        if p:
            politica_global = p
            break
    print(f"  Politica ITBIS para Word: {politica_global}")

    for i, item in enumerate(items):
        if i >= len(filas_datos):
            break
        celdas = filas_datos[i].cells

        def set_cell(col, val, celdas=celdas):
            try:
                p = celdas[col].paragraphs[0]
                p.clear()
                run = p.add_run(str(val) if val is not None else '')
                run.font.size = Pt(9)
            except Exception as e:
                print(f"Error en celda {col}: {e}")

        lote = str(item.get('lote', '')).strip()
        num = item.get('numero', i + 1)
        etiqueta_num = f"L-{lote}-{num}" if lote else str(num)

        set_cell(0, etiqueta_num)
        set_cell(1, item.get('descripcion', ''))
        set_cell(2, item.get('unidad', ''))
        set_cell(3, item.get('cantidad', ''))

        if len(celdas) > 5:
            itbis_aplica = item.get('itbis_aplica', True)
            if politica_global == 'EXENTO':
                valor_itbis = '0%'
            elif politica_global == 'INCLUIDO':
                valor_itbis = 'Incluido'
            elif politica_global == 'TRANSPARENTADO':
                valor_itbis = '18%' if itbis_aplica else '0%'
            else:
                valor_itbis = '18%' if itbis_aplica else '0%'
            set_cell(5, valor_itbis)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


@app.route('/agente-033', methods=['POST'])
def agente_033():
    try:
        data = request.get_json()
        referencia = data.get('referencia')

        if not referencia:
            return jsonify({"error": "Falta 'referencia'"}), 400

        print(f"\nAGENTE 033: {referencia}")
        nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)
        zip_path = f"{TEMP_DIR}/{nombre_seguro}.zip"

        print("PASO 1: Descargando ZIP...")
        os.makedirs(TEMP_DIR, exist_ok=True)

        zip_ya_existe = os.path.exists(zip_path) and \
            (time.time() - os.path.getmtime(zip_path)) / 86400 <= CACHE_DIAS

        if not zip_ya_existe:
            descargar_pliego(referencia, guardar_zip=True)
        else:
            print("ZIP en cache")

        if not os.path.exists(zip_path):
            return jsonify({"error": "No se pudo obtener el ZIP"}), 500

        print("PASO 2: Extrayendo archivos...")
        f033_bytes = None
        fichas_prioritarias = []
        fichas_pliego = []
        fichas_secundarias = []

        try:
            zf_handle = zipfile.ZipFile(zip_path, 'r')
        except zipfile.BadZipFile:
            return jsonify({
                "error": "El archivo descargado del SECP esta corrupto o incompleto. Intenta de nuevo."
            }), 500
        except Exception as e:
            return jsonify({
                "error": f"No se pudo abrir el archivo ZIP: {str(e)}"
            }), 500

        with zf_handle as zf:
            archivos = zf.namelist()

            if not archivos:
                return jsonify({"error": "El ZIP descargado esta vacio."}), 500

            tiene_adjuntos = any('1_Publicaciones/Adjuntos/' in a for a in archivos)
            if not tiene_adjuntos:
                return jsonify({
                    "error": "El ZIP no contiene la carpeta 1_Publicaciones/Adjuntos/. La estructura del expediente es diferente a la esperada."
                }), 500

            nombre_ficha_encontrada = None
            for archivo in archivos:
                if '1_Publicaciones/Adjuntos/' not in archivo:
                    continue
                nombre = os.path.basename(archivo).lower()

                if archivo.lower().endswith(('.docx', '.doc')):
                    if '033' in nombre:
                        f033_bytes = zf.read(archivo)
                        print(f"  F033 encontrado: {os.path.basename(archivo)}")

                if archivo.lower().endswith('.pdf'):
                    es_ficha   = any(k in nombre for k in ['ficha', 'tecnica'])
                    es_pliego  = any(k in nombre for k in ['pliego', 'condiciones', 'terminos'])
                    es_listado = any(k in nombre for k in ['listado', 'especificacion'])
                    if es_ficha:
                        fichas_prioritarias.append(zf.read(archivo))
                        if not nombre_ficha_encontrada:
                            nombre_ficha_encontrada = os.path.basename(archivo)
                        print(f"  Ficha tecnica: {os.path.basename(archivo)}")
                    elif es_pliego:
                        fichas_pliego.append(zf.read(archivo))
                        print(f"  Pliego: {os.path.basename(archivo)}")
                    elif es_listado:
                        fichas_secundarias.append(zf.read(archivo))
                        print(f"  Listado: {os.path.basename(archivo)}")

        if not f033_bytes:
            try:
                with zipfile.ZipFile(zip_path, 'r') as zf_retry:
                    for nombre_doc, doc_bytes in extraer_docx_de_adjuntos(zf_retry):
                        if '033' in nombre_doc:
                            f033_bytes = doc_bytes
                            print(f"  F033 encontrado en ZIP anidado: {nombre_doc}")
                            break
            except Exception as e:
                print(f"  Error buscando F033 en ZIP anidado: {e}")

        usar_ficha_fallback = False
        if not f033_bytes:
            if fichas_prioritarias:
                usar_ficha_fallback = True
                print(f"  F033 no encontrado — usando ficha tecnica como fallback")
            else:
                return jsonify({
                    "error": "No se encontro el F033 (.docx) en 1_Publicaciones/Adjuntos/ ni en ZIPs anidados. Esta licitacion puede ser Comparacion de Precios."
                }), 404

        candidatos = []
        if fichas_pliego:
            candidatos.append(('pliego', fichas_pliego))
        if fichas_prioritarias:
            candidatos.append(('ficha tecnica', fichas_prioritarias))
        if fichas_secundarias:
            candidatos.append(('listado', fichas_secundarias))

        if not candidatos:
            return jsonify({
                "error": "No es posible procesar el formulario porque no se encontraron PDFs con items en 1_Publicaciones/Adjuntos/."
            }), 404

        print(f"  F033: {'SI' if not usar_ficha_fallback else 'NO (usando ficha tecnica)'} | Candidatos: {[c[0] for c in candidatos]}")

        print("PASO 3: Extrayendo items con Claude...")
        items = []
        for nombre_candidato, fichas_bytes in candidatos:
            print(f"  Probando con: {nombre_candidato}")
            items = extraer_items_con_claude(fichas_bytes, referencia)
            print(f"  -> {len(items)} items encontrados")
            if len(items) >= 5:
                print(f"  Usando {nombre_candidato} ({len(items)} items)")
                break
            else:
                print(f"  Muy pocos items en {nombre_candidato}, probando siguiente...")

        if not items:
            return jsonify({"error": "Claude no extrajo items de ninguno de los PDFs disponibles"}), 500

        print("PASO 4: Generando formulario de oferta...")
        if usar_ficha_fallback:
            docx_relleno = generar_oferta_desde_ficha(
                items,
                nombre_ficha=nombre_ficha_encontrada or "ficha tecnica"
            )
            nombre_descarga = f"Ficha_Tecnica_{nombre_seguro}.docx"
            print(f"  Oferta desde ficha tecnica lista ({len(docx_relleno)} bytes)")
        else:
            docx_relleno = llenar_f033(f033_bytes, items)
            nombre_descarga = f"F033_{nombre_seguro}.docx"
            print(f"  F033 pre-llenado listo ({len(docx_relleno)} bytes)")

        return send_file(
            io.BytesIO(docx_relleno),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=nombre_descarga
        )

    except Exception as e:
        print(f"Agente 033 error: {str(e)}")
        return jsonify({"error": str(e)}), 500
_descarga_estado = {
    "corriendo": False,
    "total": 0,
    "procesadas": 0,
    "pdfs_guardados": 0,
    "errores": 0,
    "ultimo_mensaje": "Sin ejecutar",
    "iniciado_en": None
}

PALABRAS_INCLUIR_OFERTA = ['oferta', 'cotiz', 'econom', 'precio', 'propuesta']
PALABRAS_EXCLUIR_OFERTA = [
    'registro', 'mercantil', 'tss', 'dgii', 'constancia',
    'tecnic', 'rpe', 'rnc', 'cedula', 'pasaporte',
    'balance', 'financ', 'declaracion'
]


def _es_oferta_economica(nombre_pdf):
    n = nombre_pdf.lower()
    if 'ranl' in n and 'resumen' not in n:
        return False
    for palabra in PALABRAS_EXCLUIR_OFERTA:
        if palabra in n:
            return False
    for palabra in PALABRAS_INCLUIR_OFERTA:
        if palabra in n:
            return True
    if n in ('oferta.pdf', 'ofertas.pdf'):
        return True
    return True


def _worker_descarga(lote_size, db_url):
    global _descarga_estado
    _descarga_estado["corriendo"] = True
    _descarga_estado["pdfs_guardados"] = 0
    _descarga_estado["errores"] = 0
    _descarga_estado["procesadas"] = 0
    _descarga_estado["iniciado_en"] = datetime.now().strftime('%H:%M:%S')

    try:
        conn = psycopg2.connect(db_url)
        cur = conn.cursor()
        cur.execute("""
            SELECT l.id, l.referencia, l.descripcion
            FROM   licitaciones l
            WHERE  l.estado = 'Proceso adjudicado y celebrado'
              AND  NOT EXISTS (
                       SELECT 1 FROM ofertas_pendientes op
                       WHERE  op.licitacion_id = l.id
                   )
              AND (
                  l.descripcion ILIKE '%%reactivo%%'
                  OR l.descripcion ILIKE '%%laboratorio%%'
                  OR l.descripcion ILIKE '%%calibr%%'
                  OR l.descripcion ILIKE '%%microbiolog%%'
                  OR l.descripcion ILIKE '%%instrumento%%'
                  OR l.descripcion ILIKE '%%quimico%%'
                  OR l.descripcion ILIKE '%%medicamento%%'
                  OR l.descripcion ILIKE '%%insumo%%'
                  OR l.descripcion ILIKE '%%material medico%%'
              )
            ORDER  BY l.id
            LIMIT  %s
        """, (lote_size,))
        pendientes = cur.fetchall()
        cur.close()
        conn.close()

        _descarga_estado["total"] = len(pendientes)
        _descarga_estado["ultimo_mensaje"] = f"{len(pendientes)} licitaciones pendientes"
        print(f"DESCARGA BACKFILL: {len(pendientes)} licitaciones")

        for (lid, referencia, descripcion) in pendientes:
            _descarga_estado["ultimo_mensaje"] = f"Descargando {referencia}..."
            print(f"\n[{lid}] {referencia}")

            try:
                nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)
                zip_path = f"{TEMP_DIR}/{nombre_seguro}.zip"
                os.makedirs(TEMP_DIR, exist_ok=True)

                zip_en_cache = (
                    os.path.exists(zip_path) and
                    (time.time() - os.path.getmtime(zip_path)) / 86400 <= CACHE_DIAS
                )
                if not zip_en_cache:
                    descargar_pliego(referencia, guardar_zip=True)

                if not os.path.exists(zip_path):
                    print(f"  ZIP no disponible")
                    _descarga_estado["errores"] += 1
                    _descarga_estado["procesadas"] += 1
                    continue

                pdfs_guardados_esta = 0
                with zipfile.ZipFile(zip_path, 'r') as zf:
                    pdfs_oferta = [
                        a for a in zf.namelist()
                        if re.search(r'3_Ofer', a, re.IGNORECASE)
                        and a.lower().endswith('.pdf')
                        and _es_oferta_economica(os.path.basename(a))
                    ]
                    print(f"  PDFs de oferta encontrados: {len(pdfs_oferta)}")

                    conn2 = psycopg2.connect(db_url)
                    cur2 = conn2.cursor()

                    for ruta_en_zip in pdfs_oferta:
                        nombre_pdf = os.path.basename(ruta_en_zip)
                        _descarga_estado["ultimo_mensaje"] = f"{referencia} -> {nombre_pdf}"
                        try:
                            pdf_bytes = zf.read(ruta_en_zip)
                            cur2.execute("""
                                INSERT INTO ofertas_pendientes
                                    (licitacion_id, referencia, nombre_procedimiento,
                                     nombre_pdf, pdf_bytes)
                                VALUES (%s, %s, %s, %s, %s)
                                ON CONFLICT (licitacion_id, nombre_pdf) DO NOTHING
                            """, (lid, referencia, descripcion, nombre_pdf,
                                  psycopg2.Binary(pdf_bytes)))
                            pdfs_guardados_esta += 1
                            print(f"  Guardado: {nombre_pdf}")
                        except Exception as e:
                            print(f"  Error guardando {nombre_pdf}: {e}")
                            conn2.rollback()

                    conn2.commit()
                    cur2.close()
                    conn2.close()

                _descarga_estado["pdfs_guardados"] += pdfs_guardados_esta

            except Exception as e:
                print(f"  Error en {referencia}: {e}")
                _descarga_estado["errores"] += 1

            _descarga_estado["procesadas"] += 1

        _descarga_estado["ultimo_mensaje"] = (
            f"Completado: {_descarga_estado['procesadas']} licitaciones, "
            f"{_descarga_estado['pdfs_guardados']} PDFs guardados, "
            f"{_descarga_estado['errores']} errores"
        )
        print(f"\nDESCARGA COMPLETADA: {_descarga_estado['ultimo_mensaje']}")

    except Exception as e:
        _descarga_estado["ultimo_mensaje"] = f"Error fatal: {str(e)}"
        print(f"DESCARGA ERROR FATAL: {e}")
    finally:
        _descarga_estado["corriendo"] = False


@app.route('/iniciar-descarga-backfill', methods=['POST'])
def iniciar_descarga_backfill():
    token   = request.headers.get('X-Backfill-Token', '')
    secreto = os.environ.get('BACKFILL_SECRET', '')
    if not secreto or token != secreto:
        return jsonify({"error": "No autorizado"}), 401

    if _descarga_estado["corriendo"]:
        return jsonify({
            "error": "Ya hay una descarga en curso",
            "estado": _descarga_estado
        }), 409

    data      = request.get_json() or {}
    lote_size = int(data.get('lote_size', 50))
    db_url    = os.environ.get('DATABASE_URL')
    if not db_url:
        return jsonify({"error": "DATABASE_URL no configurada"}), 500

    hilo = threading.Thread(
        target=_worker_descarga,
        args=(lote_size, db_url),
        daemon=True
    )
    hilo.start()

    return jsonify({
        "status": "iniciado",
        "lote_size": lote_size,
        "mensaje": f"Descargando hasta {lote_size} licitaciones en background"
    })


@app.route('/descarga-backfill-status', methods=['GET'])
def descarga_backfill_status():
    token   = request.headers.get('X-Backfill-Token', '')
    secreto = os.environ.get('BACKFILL_SECRET', '')
    if not secreto or token != secreto:
        return jsonify({"error": "No autorizado"}), 401
    return jsonify(_descarga_estado)


@app.route('/organizador-oferta', methods=['POST'])
def organizador_oferta():
    try:
        data          = request.get_json()
        empresa_id    = data.get('empresa_id')
        referencia    = data.get('referencia')
        licitacion    = data.get('licitacion', {})
        dictamen      = data.get('dictamen', {})

        if not referencia or not licitacion or not dictamen:
            return jsonify({'success': False, 'error': 'referencia, licitacion y dictamen son requeridos'}), 400

        empresa_desc = ''
        db_url = os.environ.get('DATABASE_URL')
        if db_url and empresa_id:
            try:
                conn = psycopg2.connect(db_url)
                cur  = conn.cursor()
                cur.execute('SELECT descripcion FROM empresas WHERE id = %s LIMIT 1', (empresa_id,))
                fila = cur.fetchone()
                if fila:
                    empresa_desc = fila[0] or ''
                cur.close()
                conn.close()
            except Exception as e:
                print(f'Error leyendo empresa: {e}')

        analisis_pliego = None
        if db_url and empresa_id:
            try:
                conn = psycopg2.connect(db_url)
                cur  = conn.cursor()
                cur.execute(
                    'SELECT analisis_json FROM analisis_pliegos WHERE empresa_id = %s AND referencia = %s LIMIT 1',
                    (empresa_id, referencia)
                )
                fila = cur.fetchone()
                if fila:
                    analisis_pliego = json.loads(fila[0])
                cur.close()
                conn.close()
            except Exception as e:
                print(f'Error leyendo analisis_pliegos: {e}')

        if analisis_pliego:
            requisitos = analisis_pliego.get('requisitos', [])[:5]
            req_texto  = '\n'.join(f'- {r}' for r in requisitos) if requisitos else '- No disponible'
            garantias  = analisis_pliego.get('viabilidad', {}).get('garantias', 'No especificado')
            experiencia = analisis_pliego.get('viabilidad', {}).get('experiencia_previa', 'No especificado')
            seccion_pliego = f"""REQUISITOS DEL PLIEGO:
{req_texto}
Garantias exigidas: {garantias}
Experiencia previa: {experiencia}"""
        else:
            seccion_pliego = 'ANALISIS DEL PLIEGO: No disponible — usar descripcion de la licitacion como referencia.'

        condiciones = dictamen.get('condiciones', [])
        condiciones_texto = '\n'.join(
            f"- {'[URGENTE] ' if c.get('urgente') else ''}{c.get('texto', '')}"
            for c in condiciones
        )

        monto_fmt = f"RD${float(licitacion.get('monto', 0)):,.0f}"

        prompt = f"""Eres el Organizador de Oferta de Compita. Genera un plan de trabajo estructurado para que una empresa licitadora participe en una licitacion publica especifica de Republica Dominicana.

El output sera pegado directamente en KanbanBonsai para generar un Bonsai con 5 sprints. Cada hoja (tarea) debe ser concreta, accionable y especifica para ESTA licitacion — no generica.

LICITACION:
- Referencia: {referencia}
- Descripcion: {licitacion.get('descripcion', '')}
- Entidad: {licitacion.get('entidad', '')}
- Tipo de proceso: {licitacion.get('tipo', '')}
- Monto estimado: {monto_fmt}
- Dias disponibles: {licitacion.get('diasDisponibles', '')}
- Fecha limite: {licitacion.get('fecha_presentacion', 'No especificada')}

EMPRESA LICITADORA:
{empresa_desc}

VEREDICTO DEL COACH: {dictamen.get('veredicto', '')}
CONDICIONES A ATENDER:
{condiciones_texto}

{seccion_pliego}

Responde UNICAMENTE con este formato, sin texto adicional ni explicaciones:

PROYECTO: [nombre del Bonsai, maximo 8 palabras, especifico para esta licitacion]
DESCRIPCION: [2 lineas sobre el objetivo de este plan de oferta]

SPRINT 1 — Analisis y Evaluacion
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]

SPRINT 2 — Documentacion Legal
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]

SPRINT 3 — Oferta Tecnica
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]

SPRINT 4 — Oferta Economica
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]

SPRINT 5 — Entrega y Seguimiento
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]
- [tarea especifica para esta licitacion]"""

        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            return jsonify({'success': False, 'error': 'ANTHROPIC_API_KEY no configurada'}), 500

        print(f'Generando plan de oferta para {referencia}...')
        response = requests.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': api_key,
                'anthropic-version': '2023-06-01'
            },
            json={
                'model': 'claude-sonnet-4-20250514',
                'max_tokens': 1200,
                'messages': [{'role': 'user', 'content': prompt}]
            },
            timeout=90
        )

        if response.status_code != 200:
            raise Exception(f'Error Claude API: {response.status_code}')

        plan_generado = response.json()['content'][0]['text'].strip()
        print(f'Plan generado: {len(plan_generado)} caracteres')

        return jsonify({'success': True, 'prompt': plan_generado})

    except Exception as e:
        print(f'Error en organizador-oferta: {str(e)}')
        return jsonify({'success': False, 'error': str(e)}), 500


REPORTES_DIR = "/tmp/reportes"
F033_DIR     = "/tmp/f033"


def verificar_f033_en_cache(referencia):
    try:
        os.makedirs(F033_DIR, exist_ok=True)
        nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)
        for prefijo in ['F033', 'Ficha_Tecnica']:
            ruta = os.path.join(F033_DIR, f"{prefijo}_{nombre_seguro}.docx")
            if os.path.exists(ruta):
                edad_dias = (time.time() - os.path.getmtime(ruta)) / 86400
                if edad_dias <= CACHE_DIAS:
                    return ruta
        return None
    except Exception as e:
        print(f"Error verificando cache F033: {e}")
        return None


def generar_f033_y_cachear(referencia):
    nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)
    zip_path = f"{TEMP_DIR}/{nombre_seguro}.zip"

    if not os.path.exists(zip_path):
        return None, "ZIP no disponible — no se puede generar formulario"

    try:
        f033_bytes = None
        fichas_prioritarias, fichas_pliego, fichas_secundarias = [], [], []
        nombre_ficha_encontrada = None

        with zipfile.ZipFile(zip_path, 'r') as zf:
            archivos = zf.namelist()
            if not any('1_Publicaciones/Adjuntos/' in a for a in archivos):
                return None, "ZIP sin carpeta de adjuntos esperada"

            for archivo in archivos:
                if '1_Publicaciones/Adjuntos/' not in archivo:
                    continue
                nombre = os.path.basename(archivo).lower()
                if archivo.lower().endswith(('.docx', '.doc')) and '033' in nombre:
                    f033_bytes = zf.read(archivo)
                if archivo.lower().endswith('.pdf'):
                    if any(k in nombre for k in ['ficha', 'tecnica']):
                        fichas_prioritarias.append(zf.read(archivo))
                        if not nombre_ficha_encontrada:
                            nombre_ficha_encontrada = os.path.basename(archivo)
                    elif any(k in nombre for k in ['pliego', 'condiciones', 'terminos']):
                        fichas_pliego.append(zf.read(archivo))
                    elif any(k in nombre for k in ['listado', 'especificacion']):
                        fichas_secundarias.append(zf.read(archivo))

        usar_ficha_fallback = False
        if not f033_bytes:
            if fichas_prioritarias:
                usar_ficha_fallback = True
                print(f"  generar_f033_y_cachear: F033 no encontrado, usando ficha tecnica")
            else:
                return None, "F033 (.docx) no encontrado y no hay ficha tecnica disponible"

        candidatos = []
        if fichas_pliego:       candidatos.append(fichas_pliego)
        if fichas_prioritarias: candidatos.append(fichas_prioritarias)
        if fichas_secundarias:  candidatos.append(fichas_secundarias)
        if not candidatos:
            return None, "No hay PDFs con items en el ZIP"

        items = []
        for fichas in candidatos:
            items = extraer_items_con_claude(fichas, referencia)
            if len(items) >= 5:
                break

        if not items:
            return None, "No se extrajeron items del ZIP"

        if usar_ficha_fallback:
            docx_bytes = generar_oferta_desde_ficha(
                items,
                nombre_ficha=nombre_ficha_encontrada or "ficha tecnica"
            )
            prefijo = "Ficha_Tecnica"
        else:
            docx_bytes = llenar_f033(f033_bytes, items)
            prefijo = "F033"

        os.makedirs(F033_DIR, exist_ok=True)
        ruta = os.path.join(F033_DIR, f"{prefijo}_{nombre_seguro}.docx")
        with open(ruta, 'wb') as f:
            f.write(docx_bytes)

        print(f"  Formulario generado y cacheado: {ruta}")
        return ruta, None

    except Exception as e:
        print(f"Error generando formulario: {e}")
        return None, str(e)


def mapear_catalogo_con_claude(empresa_desc, requisitos, api_key):
    if not empresa_desc or not requisitos:
        return None
    prompt = f"""Analiza si esta empresa puede suplir los requisitos de una licitacion publica.

EMPRESA:
{empresa_desc}

REQUISITOS DE LA LICITACION:
{chr(10).join(f'- {r}' for r in requisitos[:10])}

Responde UNICAMENTE con JSON valido, sin texto adicional:
{{
  "compatibles": ["item que la empresa puede suplir directamente"],
  "requieren_proveedor": ["item que necesita proveedor externo"],
  "sin_informacion_suficiente": false
}}

REGLA: Si la descripcion de la empresa no tiene suficiente detalle,
pon sin_informacion_suficiente: true y listas vacias. No inventes."""
    try:
        resp = requests.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': api_key,
                'anthropic-version': '2023-06-01'
            },
            json={
                'model': 'claude-sonnet-4-20250514',
                'max_tokens': 800,
                'messages': [{'role': 'user', 'content': prompt}]
            },
            timeout=30
        )
        if resp.status_code != 200:
            return None
        texto = resp.json()['content'][0]['text'].strip()
        texto = texto.replace('```json', '').replace('```', '').strip()
        inicio = texto.find('{')
        fin = texto.rfind('}')
        if inicio == -1 or fin == -1:
            return None
        return json.loads(texto[inicio:fin+1])
    except Exception as e:
        print(f'Error en mapeo catalogo: {e}')
        return None


def obtener_estado_perfil_licitador(empresa_id, db_url):
    if not db_url or not empresa_id:
        return None
    try:
        conn = psycopg2.connect(db_url)
        cur = conn.cursor()
        cur.execute("""
            SELECT nombre, es_permanente, fecha_vencimiento
            FROM perfil_licitador
            WHERE empresa_id = %s
            ORDER BY grupo ASC, orden ASC
        """, (empresa_id,))
        docs = cur.fetchall()
        cur.close()
        conn.close()
        if not docs:
            return None
        hoy = datetime.now().date()
        vigentes, por_vencer, vencidos, sin_fecha, permanentes = [], [], [], [], []
        for nombre, es_permanente, fecha_venc in docs:
            if es_permanente:
                permanentes.append(nombre)
                continue
            if not fecha_venc:
                sin_fecha.append(nombre)
                continue
            if hasattr(fecha_venc, 'date'):
                fecha_venc = fecha_venc.date()
            diff = (fecha_venc - hoy).days
            if diff < 0:
                vencidos.append(f"{nombre} (vencio hace {abs(diff)} dias)")
            elif diff <= 30:
                por_vencer.append(f"{nombre} (vence en {diff} dias)")
            else:
                vigentes.append(nombre)
        return {
            'permanentes': permanentes,
            'vigentes': vigentes,
            'por_vencer': por_vencer,
            'vencidos': vencidos,
            'sin_fecha': sin_fecha,
            'total': len(docs)
        }
    except Exception as e:
        print(f'Error obteniendo perfil licitador: {e}')
        return None


def analizar_pliego_desde_cache(pdf_path, referencia, titulo, empresa_desc, api_key, monto=0):
    try:
        es_pdf_imagen = False
        texto_completo = ""
        try:
            reader = PdfReader(pdf_path)
            for pg in reader.pages:
                try:
                    t = pg.extract_text()
                    if t:
                        texto_completo += t + "\n\n"
                except Exception:
                    continue
        except Exception as e:
            print(f"Error leyendo PDF: {e}")

        if len(texto_completo.strip()) < 200:
            print("analizar_pliego_desde_cache: PDF imagen detectado — usando envio nativo a Claude")
            es_pdf_imagen = True
            texto_completo = "[El pliego es un PDF de imagen — se adjunta como documento para analisis directo por Claude]"
        elif len(texto_completo) > 100000:
            texto_completo = texto_completo[:100000] + "\n\n[DOCUMENTO TRUNCADO]"

        fecha_hoy = datetime.now().strftime('%d/%m/%Y')
        perfil_txt = f"\nDescripcion: {empresa_desc}" if empresa_desc else ""
        monto_txt  = f"\n- Monto oficial: RD${float(monto):,.2f} — usa esta cifra en la sintesis, no extraigas un monto diferente del documento" if monto else ""
        prompt = f"""Eres un experto analista de licitaciones publicas dominicanas.

CONTEXTO:
- Referencia: {referencia}
- Titulo: {titulo}{monto_txt}
{perfil_txt}

---INICIO DEL PLIEGO---
{texto_completo}
---FIN DEL PLIEGO---

Analiza el pliego y responde SOLO con este JSON:
{{
  "sintesis": "Resumen ejecutivo en 2-3 oraciones",
  "oportunidades": ["oportunidad 1", "oportunidad 2"],
  "riesgos": ["riesgo 1", "riesgo 2"],
  "requisitos": ["requisito 1", "requisito 2", "requisito 3"],
  "certificaciones_iso": {{
    "exige_iso": "SI o NO",
    "listado": [],
    "nota": ""
  }},
  "tiempos": {{
    "fecha_limite_oferta": "DD/MM/YYYY",
    "dias_calendario_restantes": "N dias",
    "alerta": "HOLGADO, AJUSTADO, o MUY AJUSTADO",
    "fechas_clave": [],
    "advertencia": ""
  }},
  "viabilidad": {{
    "veredicto": "VIABLE, VIABLE CON RIESGOS, o DIFICIL DE CUMPLIR",
    "garantias": "descripcion de garantias",
    "experiencia_previa": "experiencia requerida",
    "especificaciones_tecnicas": "compatibilidad"
  }},
  "evaluacion": {{
    "a_favor": [],
    "en_contra": []
  }},
  "precios_historicos": []
}}"""

        if es_pdf_imagen:
            try:
                tamano_pdf_mb = os.path.getsize(pdf_path) / (1024 * 1024)
            except Exception:
                tamano_pdf_mb = 0
            if tamano_pdf_mb > 20:
                print(f"analizar_pliego_desde_cache: PDF imagen de {tamano_pdf_mb:.1f} MB — demasiado grande, omitiendo")
                return None
            print("analizar_pliego_desde_cache: enviando PDF como documento nativo...")
            with open(pdf_path, 'rb') as f:
                pdf_b64 = base64.b64encode(f.read()).decode('utf-8')
            messages = [{'role': 'user', 'content': [
                {'type': 'document', 'source': {'type': 'base64', 'media_type': 'application/pdf', 'data': pdf_b64}},
                {'type': 'text', 'text': prompt}
            ]}]
        else:
            messages = [{'role': 'user', 'content': prompt}]

        resp = requests.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': api_key,
                'anthropic-version': '2023-06-01'
            },
            json={
                'model': 'claude-sonnet-4-20250514',
                'max_tokens': 3000,
                'messages': messages
            },
            timeout=90
        )
        if resp.status_code == 413:
            print(f"analizar_pliego_desde_cache: PDF demasiado grande para Claude API — omitiendo analisis")
            return None
        if resp.status_code != 200:
            return None
        texto_resp = resp.json()['content'][0]['text']
        texto_resp = texto_resp.replace('```json', '').replace('```', '').strip()
        inicio = texto_resp.find('{')
        fin = texto_resp.rfind('}')
        if inicio == -1 or fin == -1:
            return None
        return json.loads(texto_resp[inicio:fin+1])
    except Exception as e:
        print(f'Error en analisis desde cache: {e}')
        return None


def generar_prompt_kanban(referencia, licitacion, dictamen, analisis_pliego, empresa_desc, api_key,
                          perfil=None, formulario_disponible=False, formulario_es_ficha=False):
    try:
        # ── Inventario de outputs de Claude para Sprint 1 ──────────────────
        outputs_claude = []
        if analisis_pliego:
            n_req    = len(analisis_pliego.get('requisitos', []))
            n_riesg  = len(analisis_pliego.get('riesgos', []))
            viab     = analisis_pliego.get('viabilidad', {}).get('veredicto', '')
            outputs_claude.append(
                f"- Analisis del pliego: {n_req} requisitos, {n_riesg} riesgos identificados, viabilidad {viab}"
            )
        if formulario_disponible:
            tipo_form = "desde ficha tecnica (sin F033 oficial)" if formulario_es_ficha else "F033 oficial pre-llenado"
            outputs_claude.append(
                f"- Formulario de oferta pre-llenado ({tipo_form}) — disponible en el Reporte para descargar"
            )
        if perfil:
            p = perfil
            outputs_claude.append(
                f"- Estado del Perfil Licitador: {len(p.get('permanentes',[]))} permanentes, "
                f"{len(p.get('vigentes',[]))} vigentes, "
                f"{len(p.get('por_vencer',[]))} por vencer, "
                f"{len(p.get('vencidos',[]))} vencidos"
            )
        if analisis_pliego and analisis_pliego.get('requisitos'):
            outputs_claude.append(
                f"- Checklist de {len(analisis_pliego['requisitos'])} documentos/requisitos de entrega extraidos del pliego"
            )
        outputs_texto = '\n'.join(outputs_claude) or '- No hay outputs adicionales de Claude'

        # ── Tareas documentales reales (solo lo que requiere accion) ───────
        doc_tipo_a = []  # renovaciones del Perfil
        doc_tipo_b = []  # requisitos especificos de la licitacion

        if perfil:
            for d in perfil.get('vencidos', []):
                doc_tipo_a.append(f"URGENTE — Renovar: {d}")
            for d in perfil.get('por_vencer', []):
                doc_tipo_a.append(f"Renovar antes del cierre: {d}")

        garantias = (analisis_pliego or {}).get('viabilidad', {}).get('garantias', '')
        if garantias:
            doc_tipo_b.append(f"Tramitar: {garantias[:120]}")

        certs = (analisis_pliego or {}).get('certificaciones_iso', {})
        if certs.get('exige_iso') == 'SI':
            for c in certs.get('listado', []):
                doc_tipo_b.append(f"Verificar u obtener: {c}")

        if doc_tipo_a:
            doc_ctx = "RENOVACIONES PENDIENTES DEL PERFIL:\n" + '\n'.join(f"- {t}" for t in doc_tipo_a) + "\n"
        else:
            doc_ctx = "PERFIL AL DIA: no hay documentos vencidos ni por vencer.\n"
        if doc_tipo_b:
            doc_ctx += "REQUISITOS ESPECIFICOS DE ESTA LICITACION (no cubiertos por el Perfil):\n"
            doc_ctx += '\n'.join(f"- {t}" for t in doc_tipo_b)
        else:
            doc_ctx += "No se detectaron requisitos documentales adicionales especificos para esta licitacion."

        # ── Requisitos tecnicos para Sprint 3 ─────────────────────────────
        req_lista = (analisis_pliego or {}).get('requisitos', [])[:6]
        req_texto = '\n'.join(f'- {r}' for r in req_lista) or '- Ver analisis del pliego'

        # ── Datos de la licitacion ─────────────────────────────────────────
        dias         = licitacion.get('diasDisponibles', '')
        fecha_limite = licitacion.get('fecha_presentacion', 'No especificada')
        monto_val    = licitacion.get('monto', 0)
        try:
            monto_fmt = f"RD${float(monto_val):,.0f}"
        except Exception:
            monto_fmt = str(monto_val)

        instruccion_form = (
            "Completar precios unitarios en el formulario pre-llenado descargado en Sprint 1 (items y cantidades ya estan)"
            if formulario_disponible else
            "Crear formulario de oferta economica con todos los items, unidades, cantidades y precios"
        )

        condiciones_texto = '\n'.join(
            f"- {'[URGENTE] ' if c.get('urgente') else ''}{c.get('texto', '')}"
            for c in dictamen.get('condiciones', [])
        )

        prompt = f"""Genera el plan de trabajo en KanbanBonsai para preparar la oferta en esta licitacion.
El usuario YA DECIDIO participar. El plan debe reflejar lo que Claude ya hizo y lo que solo el usuario puede hacer.

LICITACION: {referencia}
Descripcion: {licitacion.get('descripcion', '')}
Entidad: {licitacion.get('entidad', '')}
Tipo: {licitacion.get('tipo', '')} | Monto: {monto_fmt}
Dias disponibles para presentar la oferta: {dias} | Fecha limite de presentacion: {fecha_limite}
Empresa: {empresa_desc}
Coach: {dictamen.get('veredicto', '')}
{condiciones_texto}

OUTPUTS QUE CLAUDE YA PREPARO:
{outputs_texto}

SITUACION DOCUMENTAL:
{doc_ctx}

ESPECIFICACIONES TECNICAS EXTRAIDAS DEL PLIEGO:
{req_texto}

INSTRUCCIONES PARA CADA SPRINT:

SPRINT 1 — Lo que Claude ya preparo:
Lista SOLO los outputs reales de la seccion OUTPUTS QUE CLAUDE YA PREPARO.
Cada tarea es una accion concreta: "Descargar X", "Revisar Y", "Anotar Z".
No agregues tareas de analisis ni de decision — Claude ya las resolvio.
3 a 4 tareas.

SPRINT 2 — Documentacion Legal:
Usa SOLO lo que dice SITUACION DOCUMENTAL.
Si el perfil esta al dia y no hay requisitos adicionales, la unica tarea es compilar el expediente.
No repitas documentos ya vigentes. 3 a 4 tareas maximo.

SPRINT 3 — Oferta Tecnica:
Tareas concretas para validar cada especificacion de ESPECIFICACIONES TECNICAS EXTRAIDAS,
conseguir fichas tecnicas de fabricantes, y preparar la propuesta tecnica. 3 a 4 tareas.

SPRINT 4 — Oferta Economica:
Primera tarea siempre: {instruccion_form}
Luego: negociar precios con proveedores, calcular totales y margenes, preparar documentos economicos requeridos.
3 a 4 tareas.

SPRINT 5 — Entrega y Seguimiento:
Ensamblar expediente con el checklist del pliego, autenticar documentos, presentar antes de {fecha_limite}
({dias} dias disponibles para preparar la oferta), obtener constancia.
NO mencionar plazos de entrega del producto — eso es post-adjudicacion.
3 a 4 tareas.

REGLA GENERAL: cada tarea es una accion ("Descargar", "Tramitar", "Confirmar", "Completar"),
nunca una evaluacion ni descripcion. No repitas lo que Claude ya hizo.

Responde UNICAMENTE con este formato:

PROYECTO: [nombre especifico maximo 8 palabras]
DESCRIPCION: [2 lineas sobre el objetivo]

SPRINT 1 — Lo que Claude ya preparo
- [tarea]
- [tarea]
- [tarea]

SPRINT 2 — Documentacion Legal
- [tarea]
- [tarea]
- [tarea]

SPRINT 3 — Oferta Tecnica
- [tarea]
- [tarea]
- [tarea]

SPRINT 4 — Oferta Economica
- [tarea]
- [tarea]
- [tarea]

SPRINT 5 — Entrega y Seguimiento
- [tarea]
- [tarea]
- [tarea]"""

        resp = requests.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': api_key,
                'anthropic-version': '2023-06-01'
            },
            json={
                'model': 'claude-sonnet-4-20250514',
                'max_tokens': 1400,
                'messages': [{'role': 'user', 'content': prompt}]
            },
            timeout=60
        )
        if resp.status_code == 200:
            return resp.json()['content'][0]['text'].strip()
        return ''
    except Exception as e:
        print(f'Error generando prompt kanban: {e}')
        return ''


def generar_html_reporte(referencia, datos):
    nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)
    licitacion    = datos.get('licitacion', {})
    dictamen      = datos.get('dictamen', {})
    analisis      = datos.get('analisis_pliego')
    perfil        = datos.get('perfil_licitador')
    mapeo         = datos.get('mapeo_catalogo')
    precios       = datos.get('precios_historicos', [])
    kanban_prompt = datos.get('kanban_prompt', '')
    f033_url      = datos.get('f033_url')
    pliego_url    = datos.get('pliego_url')
    formularios_adicionales = datos.get('formularios_adicionales', [])
    fecha_generado = datetime.now().strftime('%d %b %Y · %H:%M')

    veredicto = dictamen.get('veredicto', 'GO')
    color_v = {'GO': '#3B6D11', 'GO_RIESGO': '#f2b56b', 'NO_GO': '#DC2626'}.get(veredicto, '#3B6D11')
    label_v = {'GO': 'GO', 'GO_RIESGO': 'GO con riesgo', 'NO_GO': 'NO GO'}.get(veredicto, veredicto)

    monto_val = licitacion.get('monto', 0)
    try:
        monto_fmt = f"RD${float(monto_val):,.0f}"
    except Exception:
        monto_fmt = str(monto_val)

    fecha_lim = licitacion.get('fecha_presentacion', '')
    fecha_lim_corta = fecha_lim[:10] if fecha_lim else '—'

    def t_claude(txt, anchor=None, link_label='ver resultado'):
        a = f' — <a href="#{anchor}" style="color:#3B6D11;font-size:11px;">{link_label}</a>' if anchor else ''
        return f'<div class="task"><span class="dot green"></span><span class="task-txt">{txt}{a}</span></div>'

    def t_human(txt):
        return f'<div class="task"><span class="dot amber"></span><span class="task-txt">{txt}</span></div>'

    def sprint_sec(num, nombre, tareas):
        nc = sum(1 for t in tareas if 'dot green' in t)
        nh = sum(1 for t in tareas if 'dot amber' in t)
        return f'''<div class="sec">
  <div class="sec-hdr">
    <div style="display:flex;align-items:center;gap:8px;">
      <span class="sprint-num">Sprint {num}</span>
      <span class="sec-ttl">{nombre}</span>
    </div>
    <div style="font-size:11px;">
      <span style="color:#3B6D11;">● {nc} Claude</span>&nbsp;
      <span style="color:#BA7517;">● {nh} por hacer</span>
    </div>
  </div>
  <div class="tasks">{"".join(tareas)}</div>
</div>'''

    total_claude = 0
    total_human  = 0

    t1 = []
    if analisis and analisis.get('sintesis'):
        t1.append(t_claude(f'Analisis del pliego {referencia} completo', 'sintesis'))
        total_claude += 1
    else:
        if pliego_url:
            t1.append(f'<div class="task"><span class="dot amber"></span><span class="task-txt">Pliego no analizado — <a href="{pliego_url}" download style="background:#EAF3DE;color:#3B6D11;text-decoration:none;margin-left:6px;display:inline-block;padding:2px 9px;border-radius:4px;font-size:11px;font-weight:600;">Descargar pliego</a></span></div>')
        else:
            t1.append(t_human('Analizar pliego desde Compita — requerido para continuar'))
        total_human += 1

    if mapeo:
        nc  = len(mapeo.get('compatibles', []))
        npr = len(mapeo.get('requieren_proveedor', []))
        if mapeo.get('sin_informacion_suficiente'):
            t1.append(t_human('Mapeo catalogo vs. items — completar descripcion de empresa en Compita'))
            total_human += 1
        else:
            t1.append(t_claude(f'Mapeo catalogo: {nc} compatibles, {npr} por confirmar', 'mapeo', 'ver tabla'))
            total_claude += 1
    else:
        t1.append(t_human('Mapeo catalogo vs. items — requiere analisis del pliego'))
        total_human += 1

    t1.append(t_human('Confirmar disponibilidad de items con proveedor externo'))
    t1.append(t_human('Validar fortaleza de candidatura y decidir participacion'))
    total_human += 2
    s1 = sprint_sec('1', 'Analisis y Evaluacion', t1)

    t2 = []
    if perfil:
        nv  = len(perfil.get('vigentes', [])) + len(perfil.get('permanentes', []))
        nve = len(perfil.get('vencidos', []))
        npv = len(perfil.get('por_vencer', []))
        res = f'{nv} vigentes'
        if nve: res += f', {nve} vencidos'
        if npv: res += f', {npv} por vencer'
        t2.append(t_claude(f'Estado del Perfil Licitador — {res}', 'perfil', 'ver reporte'))
        total_claude += 1
    else:
        t2.append(t_human('Completar Perfil Licitador en Compita'))
        total_human += 1

    if analisis and analisis.get('certificaciones_iso'):
        certs = analisis['certificaciones_iso']
        if certs.get('exige_iso') == 'SI':
            lista_c = ', '.join(certs.get('listado', [])) or 'ver detalle'
            t2.append(t_claude(f'Certificaciones exigidas: {lista_c}', 'certs', 'ver analisis'))
        else:
            t2.append(t_claude('Certificaciones exigidas: ninguna ISO requerida', 'certs', 'ver analisis'))
        total_claude += 1
    else:
        t2.append(t_human('Verificar certificaciones exigidas en el pliego'))
        total_human += 1

    if perfil and perfil.get('vencidos'):
        t2.append(t_human(f'Renovar {len(perfil["vencidos"])} documento(s) vencido(s)'))
        total_human += 1
    if perfil and perfil.get('por_vencer'):
        t2.append(t_human(f'Gestionar {len(perfil["por_vencer"])} documento(s) por vencer'))
        total_human += 1

    garantias_txt = (analisis or {}).get('viabilidad', {}).get('garantias', '')
    t2.append(t_human(f'Tramitar garantia: {garantias_txt[:70]}' if garantias_txt else 'Tramitar garantia de seriedad'))
    t2.append(t_human('Compilar y verificar expediente legal completo'))
    total_human += 2
    for form in [f for f in formularios_adicionales if f.get('sprint') == 2]:
        t2.append(t_human(f"Completar {form['nombre']}"))
        total_human += 1
    s2 = sprint_sec('2', 'Documentacion Legal', t2)

    t3 = []
    if analisis and analisis.get('requisitos'):
        t3.append(t_claude(f'{len(analisis["requisitos"])} especificaciones tecnicas extraidas y cruzadas', 'checklist', 'ver tabla'))
        total_claude += 1
    else:
        t3.append(t_human('Extraer especificaciones tecnicas del pliego'))
        total_human += 1

    t3.append(t_human('Confirmar cumplimiento tecnico producto a producto'))
    t3.append(t_human('Reunir fichas tecnicas de fabricantes para items requeridos'))
    t3.append(t_human('Preparar y aprobar propuesta tecnica'))
    total_human += 3
    for form in [f for f in formularios_adicionales if f.get('sprint') == 3]:
        t3.append(t_human(f"Completar {form['nombre']}"))
        total_human += 1
    s3 = sprint_sec('3', 'Oferta Tecnica', t3)

    tiene_f033 = datos.get('tiene_f033')
    t4 = []
    if f033_url:
        label_f033 = 'F033 pre-llenado' if 'F033_' in (f033_url or '') else 'Formulario de oferta (desde ficha tecnica)'
        t4.append(f'<div class="task"><span class="dot green"></span><span class="task-txt">{label_f033} generado por Agente 033 — <a href="{f033_url}" target="_blank" style="background:#3B6D11;color:#fff;text-decoration:none;margin-left:6px;display:inline-block;padding:2px 9px;border-radius:4px;font-size:11px;font-weight:600;">Descargar</a></span></div>')
        total_claude += 1
    elif tiene_f033 is False:
        t4.append(t_human('Oferta economica manual — esta licitacion no incluye formulario F033'))
        total_human += 1
    else:
        btn_f033 = f'<button onclick="generarF033Reporte(this)" style="display:inline-block;padding:2px 9px;background:#BA7517;color:#fff;border:none;border-radius:4px;font-size:11px;font-weight:600;cursor:pointer;margin-left:6px;vertical-align:middle;">Generar formulario</button>'
        t4.append(f'<div class="task"><span class="dot amber"></span><span class="task-txt">Formulario de oferta pre-llenado{btn_f033}</span></div>')
        total_human += 1

    if precios:
        prom_global = sum(p.get('precio_promedio') or 0 for p in precios[:3]) / min(len(precios), 3)
        t4.append(t_claude(f'Precios historicos de referencia — promedio RD${prom_global:,.0f}', 'precios', 'ver tabla'))
        total_claude += 1
    else:
        t4.append(t_human('Sin precios historicos en Compita — definir precios desde cero'))
        total_human += 1

    if analisis and analisis.get('viabilidad'):
        t4.append(t_claude('Politica de ITBIS del pliego identificada'))
        total_claude += 1
    else:
        t4.append(t_human('Identificar politica de ITBIS en el pliego'))
        total_human += 1

    t4.append(t_human('Definir precios finales item por item en el formulario'))
    t4.append(t_human('Aprobar monto total y margen antes de presentar'))
    total_human += 2
    s4 = sprint_sec('4', 'Oferta Economica', t4)

    t5 = []
    if analisis and analisis.get('requisitos'):
        t5.append(t_claude(f'Checklist de entrega generado — {len(analisis["requisitos"])} documentos', 'checklist', 'ver checklist'))
        total_claude += 1
    else:
        t5.append(t_human('Generar checklist de entrega desde el pliego'))
        total_human += 1

    fecha_pliego = (analisis or {}).get('tiempos', {}).get('fecha_limite_oferta', '')
    fecha_show = fecha_pliego or fecha_lim_corta
    if fecha_show and fecha_show != '—':
        t5.append(t_claude(f'Fecha y lugar de entrega confirmados: {fecha_show}'))
        total_claude += 1
    else:
        t5.append(t_human('Confirmar fecha y lugar de entrega desde el pliego'))
        total_human += 1

    t5.append(t_human('Firmar y autenticar documentos requeridos'))
    t5.append(t_human(f'Ensamblar y entregar expediente antes del {fecha_show}'))
    t5.append(t_human('Obtener constancia de recepcion'))
    total_human += 3
    s5 = sprint_sec('5', 'Entrega y Seguimiento', t5)

    det = ''

    if analisis and analisis.get('sintesis'):
        riesgos_li = ''.join(f'<li>{r}</li>' for r in (analisis.get('riesgos') or [])[:3])
        opcs_li    = ''.join(f'<li>{o}</li>' for o in (analisis.get('oportunidades') or [])[:3])
        grid_eval  = ''
        if opcs_li or riesgos_li:
            grid_eval = f"<div class='grid2'><div class='card-green'><strong>Oportunidades</strong><ul class='ul-items'>{opcs_li}</ul></div><div class='card-amber'><strong>Riesgos</strong><ul class='ul-items'>{riesgos_li}</ul></div></div>"
        det += f'<div class="sec" id="sintesis"><div class="sec-hdr"><span class="sec-ttl">Sintesis del pliego</span><span class="bdg-ok">Claude</span></div><p class="sec-txt">{analisis["sintesis"]}</p>{grid_eval}</div>'

    if mapeo and not mapeo.get('sin_informacion_suficiente'):
        compat = mapeo.get('compatibles', [])
        prov   = mapeo.get('requieren_proveedor', [])
        li_c   = ''.join(f'<li>{i}</li>' for i in compat) or '<li>No determinado con los datos disponibles</li>'
        li_p   = ''.join(f'<li>{i}</li>' for i in prov)
        card_p = f'<div class="card-amber"><strong>{len(prov)} por confirmar con proveedor</strong><ul class="ul-items">{li_p}</ul></div>' if prov else ''
        det += f'''<div class="sec" id="mapeo">
  <div class="sec-hdr"><span class="sec-ttl">Mapeo catalogo vs. items requeridos</span><span class="bdg-ok">Claude</span></div>
  <div class="grid2">
    <div class="card-green"><strong>{len(compat)} compatibles</strong><ul class="ul-items">{li_c}</ul></div>
    {card_p}
  </div>
</div>'''

    if perfil:
        filas_p = ''
        for d in perfil['permanentes']:
            filas_p += f'<tr><td>{d}</td><td class="t-green">Permanente</td></tr>'
        for d in perfil['vigentes']:
            filas_p += f'<tr><td>{d}</td><td class="t-green">Vigente</td></tr>'
        for d in perfil['por_vencer']:
            filas_p += f'<tr><td>{d}</td><td class="t-amber">Por vencer</td></tr>'
        for d in perfil['vencidos']:
            filas_p += f'<tr><td>{d}</td><td class="t-red">Vencido</td></tr>'
        for d in perfil['sin_fecha']:
            filas_p += f'<tr><td>{d}</td><td class="t-gray">Sin fecha</td></tr>'
        det += f'<div class="sec" id="perfil"><div class="sec-hdr"><span class="sec-ttl">Perfil Licitador — detalle</span><span class="bdg-ok">Claude</span></div><table class="tbl"><tbody>{filas_p}</tbody></table></div>'

    if precios:
        filas_pr = ''
        for p in precios[:8]:
            desc = (p.get('descripcion') or '')[:55]
            prom = p.get('precio_promedio')
            pmin = p.get('precio_minimo')
            pmax = p.get('precio_maximo')
            refs = p.get('num_referencias', 0)
            mon  = p.get('moneda', 'DOP')
            if prom:
                filas_pr += f'<tr><td>{desc}</td><td>{prom:,.2f} {mon}</td><td>{pmin:,.2f}–{pmax:,.2f}</td><td>{refs}</td></tr>'
        det += f'''<div class="sec" id="precios">
  <div class="sec-hdr"><span class="sec-ttl">Precios historicos de referencia</span><span class="bdg-ok">Claude</span></div>
  <p class="sec-txt muted" style="padding-bottom:0;">Licitaciones adjudicadas en Compita. Solo referencia — no son precios finales.</p>
  <table class="tbl"><thead><tr><th>Descripcion</th><th>Promedio</th><th>Rango</th><th>Refs.</th></tr></thead><tbody>{filas_pr}</tbody></table>
</div>'''

    if analisis and analisis.get('requisitos'):
        items_li = ''.join(f'<li>{r}</li>' for r in analisis['requisitos'][:10])
        det += f'<div class="sec" id="checklist"><div class="sec-hdr"><span class="sec-ttl">Checklist de requisitos del pliego</span><span class="bdg-ok">Claude</span></div><ul class="checklist">{items_li}</ul></div>'

    if analisis and analisis.get('certificaciones_iso'):
        certs = analisis['certificaciones_iso']
        exige = certs.get('exige_iso', 'NO')
        listado_c = certs.get('listado', [])
        nota_c = certs.get('nota', '')
        if exige == 'SI':
            cuerpo_c = '<ul class="checklist">' + ''.join(f'<li>{c}</li>' for c in listado_c) + '</ul>'
        else:
            cuerpo_c = f'<p class="sec-txt">El pliego <strong>no exige</strong> certificaciones ISO ni normas equivalentes para este proceso.</p>'
        if nota_c:
            cuerpo_c += f'<p class="sec-txt muted" style="padding-top:0;">{nota_c}</p>'
        det += f'<div class="sec" id="certs"><div class="sec-hdr"><span class="sec-ttl">Certificaciones exigidas — evidencia del pliego</span><span class="bdg-ok">Claude</span></div>{cuerpo_c}</div>'

    if analisis and analisis.get('tiempos'):
        tiempos = analisis['tiempos']
        filas_t = ''
        if tiempos.get('fecha_limite_oferta'):
            filas_t += f'<tr><td>Fecha limite de oferta</td><td><strong>{tiempos["fecha_limite_oferta"]}</strong></td></tr>'
        if tiempos.get('dias_calendario_restantes'):
            filas_t += f'<tr><td>Dias restantes</td><td>{tiempos["dias_calendario_restantes"]}</td></tr>'
        alerta = tiempos.get('alerta', '')
        color_alerta = {'HOLGADO': '#3B6D11', 'AJUSTADO': '#92400E', 'MUY AJUSTADO': '#DC2626'}.get(alerta, '#6B7280')
        if alerta:
            filas_t += f'<tr><td>Estado del plazo</td><td style="color:{color_alerta};font-weight:600;">{alerta}</td></tr>'
        for fc in (tiempos.get('fechas_clave') or [])[:4]:
            filas_t += f'<tr><td>Fecha clave</td><td>{fc}</td></tr>'
        if tiempos.get('advertencia'):
            filas_t += f'<tr><td colspan="2" style="color:#92400E;font-style:italic;">{tiempos["advertencia"]}</td></tr>'
        if filas_t:
            det += f'<div class="sec" id="tiempos"><div class="sec-hdr"><span class="sec-ttl">Fechas y plazos — evidencia del pliego</span><span class="bdg-ok">Claude</span></div><table class="tbl"><tbody>{filas_t}</tbody></table></div>'

    kanban_json = json.dumps(kanban_prompt)

    return f'''<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Reporte Compita — {referencia}</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;background:#F3F4F6;color:#111827;padding:1.5rem 0}}
.wrap{{max-width:800px;margin:0 auto;padding:0 1rem}}
.hdr{{background:#3B6D11;border-radius:12px;padding:1.25rem 1.5rem;margin-bottom:1rem}}
.hdr-top{{display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:.75rem;margin-bottom:1rem}}
.hdr-label{{font-size:11px;color:#C0DD97;text-transform:uppercase;letter-spacing:.06em;margin-bottom:3px}}
.hdr-title{{font-size:17px;font-weight:600;color:#EAF3DE;margin-bottom:3px}}
.hdr-sub{{font-size:13px;color:#97C459}}
.hdr-btns{{display:flex;gap:8px}}
.btn-hdr{{padding:6px 12px;font-size:12px;background:#27500A;color:#C0DD97;border:1px solid #639922;border-radius:7px;cursor:pointer;text-decoration:none;white-space:nowrap}}
.stats{{display:grid;grid-template-columns:repeat(auto-fit,minmax(110px,1fr));gap:8px}}
.stat{{background:#27500A;border-radius:8px;padding:9px 11px}}
.stat-l{{font-size:11px;color:#97C459;margin-bottom:2px}}
.stat-v{{font-size:13px;font-weight:600;color:#EAF3DE}}
.stat-v.amber{{color:#FBBF24}}
.sec{{background:#fff;border-radius:12px;border:1px solid #E5E7EB;overflow:hidden;margin-bottom:.75rem}}
.sec-hdr{{display:flex;justify-content:space-between;align-items:center;padding:9px 14px;background:#F9FAFB;border-bottom:1px solid #E5E7EB;flex-wrap:wrap;gap:6px}}
.sprint-num{{font-size:11px;font-weight:600;background:#EAF3DE;color:#3B6D11;padding:2px 8px;border-radius:6px;white-space:nowrap}}
.sec-ttl{{font-size:13px;font-weight:500}}
.bdg-ok{{font-size:11px;background:#EAF3DE;color:#3B6D11;padding:2px 8px;border-radius:6px}}
.tasks{{padding:10px 14px;display:flex;flex-direction:column;gap:7px}}
.task{{display:flex;gap:10px;align-items:flex-start}}
.dot{{width:8px;height:8px;border-radius:50%;flex-shrink:0;margin-top:4px}}
.dot.green{{background:#3B6D11}}
.dot.amber{{background:#BA7517}}
.task-txt{{font-size:13px;color:#374151;line-height:1.5}}
.task-txt a{{text-decoration:none;margin-left:6px;vertical-align:middle;display:inline-block;padding:2px 9px;border-radius:4px;font-size:11px;font-weight:600;}}
.sec-txt{{padding:11px 14px;font-size:13px;color:#374151;line-height:1.6}}
.muted{{color:#9CA3AF!important;font-style:italic}}
.grid2{{display:grid;grid-template-columns:1fr 1fr;gap:8px;padding:11px 14px}}
.card-green{{background:#EAF3DE;border-radius:8px;padding:10px;font-size:12px;color:#3B6D11}}
.card-amber{{background:#FEF3C7;border-radius:8px;padding:10px;font-size:12px;color:#92400E}}
.ul-items{{margin:5px 0 0 14px;line-height:1.8}}
.tbl{{width:100%;font-size:13px;border-collapse:collapse}}
.tbl td,.tbl th{{padding:7px 14px;border-bottom:1px solid #F3F4F6}}
.tbl th{{background:#F9FAFB;font-size:11px;color:#6B7280;text-align:left}}
.tbl tr:last-child td{{border-bottom:none}}
.t-green{{color:#3B6D11;font-weight:500}}
.t-amber{{color:#92400E;font-weight:500}}
.t-red{{color:#DC2626;font-weight:500}}
.t-gray{{color:#9CA3AF}}
.checklist{{padding:11px 14px 11px 30px;font-size:13px;color:#374151;line-height:2.2}}
.legend{{display:flex;gap:16px;font-size:12px;color:#6B7280;margin-bottom:.75rem;padding-left:4px}}
.footer{{background:#fff;border-radius:12px;border:2px solid #3B6D11;padding:1rem 1.25rem;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:.75rem;margin-bottom:.75rem}}
.footer-t{{font-size:13px;font-weight:500}}
.footer-s{{font-size:12px;color:#6B7280;margin-top:2px}}
.btn-kb{{padding:9px 18px;background:#3B6D11;color:#fff;border:none;border-radius:8px;font-size:13px;font-weight:600;cursor:pointer}}
.divider{{font-size:11px;font-weight:600;color:#6B7280;text-transform:uppercase;letter-spacing:.06em;margin:1rem 0 .5rem;padding-left:4px}}
.meta{{font-size:11px;color:#9CA3AF;text-align:center;padding-bottom:2rem}}
@media print{{.btn-hdr,.btn-kb{{display:none}}body{{background:#fff}}}}
</style>
</head>
<body>
<div class="wrap">

<div class="hdr">
  <div class="hdr-top">
    <div>
      <div class="hdr-label">Reporte Compita</div>
      <div class="hdr-title">{referencia}</div>
      <div class="hdr-sub">{licitacion.get("entidad","")[:60]} · {licitacion.get("descripcion","")[:50]} · {monto_fmt}</div>
    </div>
    <div class="hdr-btns">
      <a href="javascript:window.print()" class="btn-hdr">PDF</a>
    </div>
  </div>
  <div class="stats">
    <div class="stat"><div class="stat-l">Veredicto Coach</div><div class="stat-v" style="color:{color_v};">{label_v}</div></div>
    <div class="stat"><div class="stat-l">Dias disponibles</div><div class="stat-v">{licitacion.get("diasDisponibles","—")}</div></div>
    <div class="stat"><div class="stat-l">Tareas Claude</div><div class="stat-v">{total_claude} listas</div></div>
    <div class="stat"><div class="stat-l">Tareas humanas</div><div class="stat-v amber">{total_human} por hacer</div></div>
  </div>
</div>

<div class="legend">
  <span><span style="display:inline-block;width:8px;height:8px;border-radius:50%;background:#3B6D11;margin-right:5px;vertical-align:middle;"></span>Claude — ya completado</span>
  <span><span style="display:inline-block;width:8px;height:8px;border-radius:50%;background:#BA7517;margin-right:5px;vertical-align:middle;"></span>Humano — por hacer en KanbanBonsai</span>
</div>

{s1}
{s2}
{s3}
{s4}
{s5}

<div class="divider">Detalle — resultados de Claude</div>
{det}

<div class="footer">
  <div>
    <div class="footer-t">Listo para trabajar en KanbanBonsai</div>
    <div class="footer-s">{total_human} tareas humanas distribuidas en 5 sprints</div>
  </div>
  <button class="btn-kb" onclick="abrirKanban()">Abrir en KanbanBonsai</button>
</div>

<div class="meta">Generado por Claude · Compita · {fecha_generado}</div>
</div>
<script>
var kp={kanban_json};
function abrirKanban(){{
  if(kp){{
    try{{navigator.clipboard.writeText(kp);}}catch(e){{}}
    var toast = document.createElement('div');
    toast.id = 'kb-toast';
    toast.style.cssText = 'position:fixed;bottom:2rem;left:50%;transform:translateX(-50%);background:#111827;color:#fff;padding:14px 22px;border-radius:10px;font-size:13px;line-height:1.6;max-width:380px;text-align:center;z-index:9999;box-shadow:0 4px 20px rgba(0,0,0,0.3);';
    toast.innerHTML = 'Plan copiado al portapapeles<br><span style="color:#9CA3AF;font-size:12px;">En KanbanBonsai haz clic en <strong style="color:#fff;">Generar con IA</strong> y pega con <strong style="color:#fff;">Ctrl+V</strong></span><br><span style="color:#6B7280;font-size:11px;margin-top:4px;display:block;">Este aviso desaparece en 60 segundos</span>';
    document.body.appendChild(toast);
    setTimeout(function(){{ window.open('https://kanban.umbusk.com/bonsais','_blank'); }}, 2000);
    setTimeout(function(){{ var t = document.getElementById('kb-toast'); if(t) t.remove(); }}, 60000);
  }} else {{
    window.open('https://kanban.umbusk.com/bonsais','_blank');
  }}
}}
async function generarF033Reporte(btn){{
  btn.disabled = true;
  btn.textContent = 'Generando... (1-2 min)';
  try{{
    var resp = await fetch('/agente-033', {{
      method: 'POST',
      headers: {{'Content-Type': 'application/json'}},
      body: JSON.stringify({{referencia: '{referencia}'}})
    }});
    if(!resp.ok) {{
      var msg = 'Error ' + resp.status;
      if(resp.status === 504) {{
        msg = 'ZIP muy grande o servidor lento. Descargalo manualmente desde el portal SECP.';
      }} else if(resp.status === 404) {{
        msg = 'No se encontro formulario (puede ser Comparacion de Precios sin formulario).';
      }} else if(resp.status === 500) {{
        try {{ var errData = await resp.json(); msg = errData.error || msg; }} catch(ex) {{}}
      }}
      throw new Error(msg);
    }}
    var blob = await resp.blob();
    var url  = URL.createObjectURL(blob);
    var a    = document.createElement('a');
    a.href   = url;
    a.download = 'Formulario_{nombre_seguro}.docx';
    a.click();
    btn.textContent = 'Descargado';
    btn.style.background = '#3B6D11';
    btn.style.color = '#fff';
  }} catch(e){{
    btn.style.background = '#DC2626';
    btn.textContent = 'Error: ' + e.message.substring(0, 60);
    setTimeout(function(){{
      btn.textContent = 'Reintentar';
      btn.style.background = '#BA7517';
      btn.disabled = false;
    }}, 5000);
  }}
}}
</script>
</body>
</html>'''


@app.route('/generar-reporte', methods=['POST'])
def generar_reporte():
    try:
        data       = request.get_json()
        empresa_id = data.get('empresa_id')
        referencia = data.get('referencia')
        licitacion = data.get('licitacion', {})
        dictamen   = data.get('dictamen', {})

        if not referencia:
            return jsonify({'success': False, 'error': 'Falta referencia'}), 400

        nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', referencia)
        os.makedirs(REPORTES_DIR, exist_ok=True)
        ruta_reporte = os.path.join(REPORTES_DIR, f"{nombre_seguro}.html")

        db_url  = os.environ.get('DATABASE_URL')
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            return jsonify({'success': False, 'error': 'ANTHROPIC_API_KEY no configurada'}), 500

        force = data.get('force', False)

        analisis_en_neon = False
        if db_url and empresa_id:
            try:
                conn = psycopg2.connect(db_url)
                cur  = conn.cursor()
                cur.execute(
                    'SELECT 1 FROM analisis_pliegos WHERE empresa_id = %s AND referencia = %s LIMIT 1',
                    (empresa_id, referencia)
                )
                analisis_en_neon = cur.fetchone() is not None
                cur.close()
                conn.close()
            except Exception:
                pass

        if not force and not analisis_en_neon:
            force = True
            print("PASO 0: Sin analisis en Neon — forzando regeneracion del reporte")

        if force and db_url and empresa_id:
            try:
                conn = psycopg2.connect(db_url)
                cur  = conn.cursor()
                cur.execute(
                    'DELETE FROM analisis_pliegos WHERE empresa_id = %s AND referencia = %s',
                    (empresa_id, referencia)
                )
                if cur.rowcount > 0:
                    print(f"PASO 0: Analisis previo borrado de Neon (force=True)")
                conn.commit()
                cur.close()
                conn.close()
            except Exception as e:
                print(f"PASO 0: Error borrando analisis de Neon: {e}")
            analisis_en_neon = False

        if not force and os.path.exists(ruta_reporte):
            edad_dias = (time.time() - os.path.getmtime(ruta_reporte)) / 86400
            if edad_dias <= CACHE_DIAS:
                dominio = os.environ.get('RAILWAY_PUBLIC_DOMAIN', 'compita-descarga-pliegos-production.up.railway.app')
                url = f"https://{dominio}/reporte/{nombre_seguro}"
                print(f"Reporte en cache: {url}")
                return jsonify({'success': True, 'url': url, 'cached': True})

        print(f"\n== PIPELINE REPORTE COMPITA: {referencia} ==")

        zip_path  = f"{TEMP_DIR}/{nombre_seguro}.zip"
        zip_ok    = os.path.exists(zip_path) and (time.time() - os.path.getmtime(zip_path)) / 86400 <= CACHE_DIAS
        pdf_cache = verificar_archivo_en_cache(referencia)

        if not zip_ok:
            if pdf_cache:
                print("PASO 1: PDF en cache — descargando ZIP para formularios y F033...")
            else:
                print("PASO 1: Descargando pliego y ZIP (primera vez)...")
            try:
                os.makedirs(TEMP_DIR, exist_ok=True)
                descargar_pliego(referencia, guardar_zip=True)
                zip_ok = os.path.exists(zip_path)
                if not pdf_cache:
                    pdf_cache = verificar_archivo_en_cache(referencia)
                print(f"  ZIP: {'OK' if zip_ok else 'no disponible'} | PDF: {'OK' if pdf_cache else 'no disponible'}")
            except Exception as e:
                print(f"  Error descargando: {e} — continuando sin ZIP/PDF")
        else:
            print(f"PASO 1: ZIP en cache OK | PDF en cache: {pdf_cache is not None}")

        formularios_adicionales = []
        tiene_f033 = None
        if zip_ok:
            try:
                FORMAS_CONOCIDAS = [
                    ({'034', 'presentofer', 'presentacion_de_oferta'}, 'F034', 'Formulario F034 — Presentacion de Oferta', 2),
                    ({'042', 'informacion_oferente', 'informacion del oferente'}, 'F042', 'Formulario F042 — Informacion del Oferente', 2),
                    ({'etico', 'integridad', 'diligencia', 'compromiso'}, 'ETICO', 'Compromiso etico de proveedores', 2),
                    ({'056', 'muestra', 'entrega_de_muestra'}, 'F056', 'Formulario F056 — Entrega de Muestras', 3),
                ]
                vistos = set()
                tiene_f033 = False
                with zipfile.ZipFile(zip_path, 'r') as zf:
                    for nombre_doc, _ in extraer_docx_de_adjuntos(zf):
                        nombre = nombre_doc.replace(' ', '_')
                        if '033' in nombre:
                            tiene_f033 = True
                            continue
                        for claves, codigo, etiqueta, sprint in FORMAS_CONOCIDAS:
                            if codigo not in vistos and any(c in nombre for c in claves):
                                formularios_adicionales.append({'codigo': codigo, 'nombre': etiqueta, 'sprint': sprint})
                                vistos.add(codigo)
                                break
                print(f"  F033 en ZIP: {'Si' if tiene_f033 else 'No'}")
                print(f"  Formularios adicionales: {[f['codigo'] for f in formularios_adicionales] or 'ninguno'}")
            except Exception as e:
                print(f"  Error inspeccionando formularios: {e}")

        empresa_desc = ''
        if db_url and empresa_id:
            try:
                conn = psycopg2.connect(db_url)
                cur  = conn.cursor()
                cur.execute('SELECT descripcion FROM empresas WHERE id = %s LIMIT 1', (empresa_id,))
                fila = cur.fetchone()
                if fila:
                    empresa_desc = fila[0] or ''
                cur.close()
                conn.close()
            except Exception as e:
                print(f'Error leyendo empresa: {e}')

        analisis_pliego = None
        if db_url and empresa_id:
            try:
                conn = psycopg2.connect(db_url)
                cur  = conn.cursor()
                cur.execute(
                    'SELECT analisis_json FROM analisis_pliegos WHERE empresa_id = %s AND referencia = %s LIMIT 1',
                    (empresa_id, referencia)
                )
                fila = cur.fetchone()
                if fila:
                    analisis_pliego = json.loads(fila[0])
                cur.close()
                conn.close()
            except Exception as e:
                print(f'Error leyendo analisis_pliegos: {e}')

        if not analisis_pliego and pdf_cache:
            print("PASO 3: Analisis del pliego no encontrado — ejecutando desde PDF...")
            titulo_lic = licitacion.get('descripcion', '')
            analisis_pliego = analizar_pliego_desde_cache(
                pdf_cache, referencia, titulo_lic, empresa_desc, api_key,
                monto=licitacion.get('monto', 0)
            )
            if analisis_pliego and db_url and empresa_id:
                try:
                    conn = psycopg2.connect(db_url)
                    cur  = conn.cursor()
                    cur.execute("""
                        INSERT INTO analisis_pliegos (empresa_id, referencia, analisis_json)
                        VALUES (%s, %s, %s)
                        ON CONFLICT (empresa_id, referencia)
                        DO UPDATE SET analisis_json = EXCLUDED.analisis_json
                    """, (empresa_id, referencia, json.dumps(analisis_pliego)))
                    conn.commit()
                    cur.close()
                    conn.close()
                    print("  Analisis guardado en Neon")
                except Exception as e:
                    print(f'  Error guardando analisis: {e}')
        elif not analisis_pliego:
            print("PASO 3: Sin PDF en cache — analisis no disponible")
        else:
            print("PASO 3: Analisis del pliego en Neon OK")

        f033_ruta = verificar_f033_en_cache(referencia)
        if not f033_ruta and zip_ok:
            print("PASO 4: Generando formulario desde ZIP...")
            f033_ruta, f033_error = generar_f033_y_cachear(referencia)
            if f033_ruta:
                print(f"  Formulario generado OK")
            else:
                print(f"  Formulario no generado: {f033_error}")
        elif f033_ruta:
            print("PASO 4: Formulario en cache OK")
        else:
            print("PASO 4: Formulario no disponible (sin ZIP)")

        dominio  = os.environ.get('RAILWAY_PUBLIC_DOMAIN', 'compita-descarga-pliegos-production.up.railway.app')
        f033_url = f"https://{dominio}/f033/{nombre_seguro}" if f033_ruta else None

        print("PASO 5: Consultando Perfil Licitador...")
        perfil = obtener_estado_perfil_licitador(empresa_id, db_url)

        print("PASO 6: Buscando precios historicos...")
        titulo_lic = licitacion.get('descripcion', '')
        precios = buscar_precios_referencia(titulo_lic, titulo_lic)
        print(f"  {len(precios)} precios encontrados")

        mapeo     = None
        requisitos = (analisis_pliego or {}).get('requisitos', [])
        if empresa_desc and requisitos:
            print("PASO 7: Mapeando catalogo vs. pliego...")
            mapeo = mapear_catalogo_con_claude(empresa_desc, requisitos, api_key)

        print("PASO 8: Generando prompt KanbanBonsai...")
        kanban_prompt = generar_prompt_kanban(
            referencia, licitacion, dictamen, analisis_pliego, empresa_desc, api_key,
            perfil=perfil,
            formulario_disponible=f033_ruta is not None,
            formulario_es_ficha=tiene_f033 is False
        )

        print("PASO 9: Generando HTML del Reporte...")
        pliego_url = f"https://{dominio}/pliego/{nombre_seguro}" if pdf_cache else None

        datos = {
            'licitacion':        licitacion,
            'dictamen':          dictamen,
            'analisis_pliego':   analisis_pliego,
            'perfil_licitador':  perfil,
            'mapeo_catalogo':    mapeo,
            'precios_historicos': precios,
            'kanban_prompt':     kanban_prompt,
            'f033_url':          f033_url,
            'pliego_url':        pliego_url,
            'formularios_adicionales': formularios_adicionales,
            'tiene_f033':              tiene_f033,
        }
        html = generar_html_reporte(referencia, datos)

        with open(ruta_reporte, 'w', encoding='utf-8') as f:
            f.write(html)

        url = f"https://{dominio}/reporte/{nombre_seguro}"
        print(f"== Reporte listo: {url} ==\n")
        return jsonify({'success': True, 'url': url, 'cached': False})

    except Exception as e:
        print(f"Error en generar-reporte: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/pliego/<nombre_seguro>', methods=['GET'])
def servir_pliego(nombre_seguro):
    nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', nombre_seguro)
    try:
        for archivo in os.listdir(TEMP_DIR):
            if archivo.startswith(nombre_seguro) and archivo.endswith('_documento.pdf'):
                ruta = os.path.join(TEMP_DIR, archivo)
                return send_file(
                    ruta,
                    mimetype='application/pdf',
                    as_attachment=True,
                    download_name=f"pliego_{nombre_seguro}.pdf"
                )
    except Exception:
        pass
    return "Pliego no encontrado en cache. Descargalo desde Compita.", 404


@app.route('/f033/<nombre_seguro>', methods=['GET'])
def servir_f033(nombre_seguro):
    nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', nombre_seguro)
    for prefijo in ['F033', 'Ficha_Tecnica']:
        ruta = os.path.join(F033_DIR, f"{prefijo}_{nombre_seguro}.docx")
        if os.path.exists(ruta):
            return send_file(
                ruta,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"{prefijo}_{nombre_seguro}.docx"
            )
    return "Formulario no encontrado. Generalo desde el Reporte Compita.", 404


@app.route('/reporte/<nombre_seguro>', methods=['GET'])
def servir_reporte(nombre_seguro):
    nombre_seguro = re.sub(r'[^a-zA-Z0-9-]', '_', nombre_seguro)
    ruta = os.path.join(REPORTES_DIR, f"{nombre_seguro}.html")
    if not os.path.exists(ruta):
        return "Reporte no encontrado. Generalo desde Compita.", 404
    return send_file(ruta, mimetype='text/html')


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)